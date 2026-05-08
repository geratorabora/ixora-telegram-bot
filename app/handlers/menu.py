# app/handlers/menu.py

from aiogram import Router, F, Bot
from aiogram.types import (
    CallbackQuery, Message, FSInputFile,
    InlineKeyboardMarkup, InlineKeyboardButton,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder
from aiogram.filters import StateFilter
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from pathlib import Path
import asyncio
import shutil
import re
import time
import sqlite3
import os
from datetime import datetime

from app.keyboards.inline_menu import (
    get_main_inline_menu,
    get_question_inline_menu,
    get_staff_inline_menu,
)
from app.logger import logger

router = Router()

# =========================
# STAFF WHITELIST
# =========================
from app.config import STAFF_USER_IDS

def is_staff(user_id: int) -> bool:
    return user_id in STAFF_USER_IDS

async def deny_staff_access(obj: Message | CallbackQuery, action: str = "unknown"):
    try:
        u = obj.from_user
        logger.warning("DENY | user_id=%s | username=%s | action=%s", u.id, u.username, action)
    except Exception:
        logger.warning("DENY | (cannot read user) | action=%s", action)
    text = "Доступ к этому разделу — только для сотрудников 🙂"
    if isinstance(obj, Message):
        await obj.answer(text)
    else:
        await obj.answer(text, show_alert=True)

def log_staff_event(user_id: int, username: str | None, action: str, allowed: bool):
    logger.info(
        "STAFF_EVENT | user_id=%s | username=%s | action=%s | allowed=%s",
        user_id, username, action, allowed,
    )

STORAGE_DIR = Path("storage")
TMP_DIR = STORAGE_DIR / "tmp"
ARCHIVE_DIR = STORAGE_DIR / "archive"
ARCHIVE_DIR.mkdir(parents=True, exist_ok=True)

# =========================
# FSM: Мастер "Объединить спецификации"
# =========================
class UploadStock(StatesGroup):
    waiting_file = State()   # ждём файл остатков от сотрудника

class MergeSpec(StatesGroup):
    waiting_pdfs = State()   # ждём PDF-ы спецификаций
    confirm     = State()    # подтверждаем / правим номер итогового документа

class LostInvoiceLetter(StatesGroup):
    waiting_company      = State()
    waiting_country      = State()
    waiting_letter_no    = State()
    waiting_letter_date  = State()
    waiting_awb          = State()
    waiting_places       = State()
    waiting_weight       = State()
    waiting_amount       = State()
    waiting_invoice_no   = State()
    waiting_invoice_date = State()
    waiting_signature    = State()

class AdjustPaymentInvoice(StatesGroup):
    waiting_file = State()


def _merge_upload_kb() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.add(InlineKeyboardButton(text="✅ Читаем файлы",  callback_data="merge:read"))
    b.add(InlineKeyboardButton(text="📎 Ещё файлы",     callback_data="merge:more"))
    b.adjust(1)
    return b.as_markup()

def _merge_confirm_kb() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.add(InlineKeyboardButton(text="🖨️ Делаем PDF", callback_data="merge:generate"))
    b.adjust(1)
    return b.as_markup()

def _merge_done_kb() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.add(InlineKeyboardButton(text="📄 Сформировать инвойс",   callback_data="merge:invoice"))
    b.add(InlineKeyboardButton(text="📦 АРХИВ",                 callback_data="merge:archive"))
    b.add(InlineKeyboardButton(text="✉️ Написать разработчику", url="tg://user?id=257207163"))
    b.adjust(1)
    return b.as_markup()

def _archive_done_kb() -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    b.add(InlineKeyboardButton(text="⬅️ Начать новое объединение", callback_data="merge:restart"))
    b.adjust(1)
    return b.as_markup()

def _lost_company_kb(companies: list[dict]) -> InlineKeyboardMarkup:
    b = InlineKeyboardBuilder()
    for c in companies[:20]:
        label = f"{c['company_name']}, {c['country']}"
        b.add(InlineKeyboardButton(
            text=label[:60],
            callback_data=f"lost_company:{c['id']}",
        ))
    b.add(InlineKeyboardButton(
        text="➕ Ввести новую компанию",
        callback_data="lost_company:new",
    ))
    b.adjust(1)
    return b.as_markup()

def _build_upload_status(spec_infos: list[str], pdf_files: list[str]) -> str:
    n_xlsx = len(spec_infos)
    n_pdf  = len(pdf_files)
    lines  = [f"📥 <b>Принято файлов: {n_xlsx + n_pdf}</b>",
              f"\n<b>XLSX (спецификации): {n_xlsx}</b>"]
    lines.extend(spec_infos)
    lines.append(f"\n<b>PDF (подписи/печати): {n_pdf}</b>")
    if pdf_files:
        for p in pdf_files:
            lines.append(f"  • {Path(p).name}")
    else:
        lines.append("  ⚠️ Не загружен — нужен для блока подписей")
    return '\n'.join(lines)


# =========================
# Мастер "Корректировка инвойса на оплату"
# =========================
PAYMENT_BANK_BLOCK_EN = [
    "IP OOO IXORA MEDICINE",
    "VAT 309915658, Reg.No 326020203217",
    "Address: Republic of Uzbekistan, 100037, Tashkent, Yakkasaray District, 1/1 Kohinur street",
    "Bank: ATB “Asia Alliance Bank”",
    "Bank address: 2a, Makhtumquli str., Yashnabad district, Tashkent, Republic of Uzbekistan",
    "SWIFT ASAC UZ 22, MFO: 01095, OKED: 64190, VAT 326020203217",
    "ac/No: 20208978105572954004 EUR",
    "ac/No: 20208840805572954004 USD",
    "ac/No: 20208643905572954004 RUB",
    "ac/No: 2020800005572954001 UZS",
    "Correspondent bank  (for payments in RUB)",
    "CB Moskommertsbank (JSC) COR/account 30111810800000058685  bic 044525951 VAT 7750005612",
]

PAYMENT_BANK_BLOCK_RU = [
    'ИП ООО "IXORA MEDICINE"',
    "ИНН 309915658, Рег.номер 326020203217",
    "Адрес: Республика Узбекистан, 100037, Ташкент, Яккасарайский район, ул. Кохинур, д. 1/1.",
    "Банк: АТБ “Азия Альянс Банк”",
    "Адрес банка: 2а, ул Махтумкули, Яшнободский район, Ташкент, Республика Узбекистан",
    "SWIFT ASAC UZ 22, МФО: 01095, ОКЭД: 64190, ИНН: 326020203217",
    "р/с: 20208978105572954004 EUR",
    "р/с: 20208840805572954004 USD",
    "р/с: 20208643905572954004 RUB",
    "р/с: 2020800005572954001 UZS",
    "Корреспондентский банк для оплат в РУБ",
    'КБ "Москоммерцбанк" (АО) кор/счет 30111810800000058685 бик 044525951 ИНН 7750005612',
]


def _adjust_payment_invoice_xlsx(src_path: Path, out_path: Path) -> tuple[bool, str]:
    from copy import copy
    from openpyxl import load_workbook
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.styles import Alignment, Border, Font

    wb = load_workbook(src_path)
    found = None

    logo_path = Path(__file__).parent.parent / "assets" / "ixora_logo_red_1.png"

    def _anchor_pos(img) -> tuple[int, int] | None:
        anchor = getattr(img, "anchor", None)
        marker = getattr(anchor, "_from", None)
        if marker is None:
            return None
        return int(marker.row) + 1, int(marker.col) + 1

    def _normalize_payment_invoice_header(ws) -> None:
        title_row = None
        for row in ws.iter_rows():
            for cell in row:
                text = str(cell.value or "").strip().lower()
                if "инвойс №" in text:
                    title_row = cell.row
                    break
            if title_row:
                break

        if not title_row:
            return

        # Верхний банковский блок из 1С не удаляем физически: в этих выгрузках много merge-областей
        # и картинок, а удаление строк легко портит drawing XML. Скрываем строки выше "Инвойс".
        for r in range(1, title_row):
            ws.row_dimensions[r].hidden = True

        # Новый логотип ставим рядом с заголовком, не трогая таблицу и существующие подписи/печати.
        has_header_logo = any(
            pos and title_row - 1 <= pos[0] <= title_row and pos[1] >= 40
            for pos in (_anchor_pos(img) for img in getattr(ws, "_images", []))
        )
        if logo_path.exists() and not has_header_logo:
            logo = XLImage(str(logo_path))
            logo.width = 105
            logo.height = 42
            ws.add_image(logo, f"AT{max(1, title_row - 1)}")

        ws.column_dimensions["A"].width = 3
        ws.column_dimensions["B"].width = 3
        ws.column_dimensions["C"].width = 3
        ws.row_dimensions[title_row].height = max(ws.row_dimensions[title_row].height or 13, 28)

    def _unmerge_intersections(ws, min_row: int, max_row: int, min_col: int, max_col: int) -> None:
        for merged in list(ws.merged_cells.ranges):
            rows_intersect = merged.min_row <= max_row and merged.max_row >= min_row
            cols_intersect = merged.min_col <= max_col and merged.max_col >= min_col
            if rows_intersect and cols_intersect:
                ws.unmerge_cells(str(merged))

    for ws in wb.worksheets:
        _normalize_payment_invoice_header(ws)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                value = str(cell.value or "").strip().lower()
                if "банковские реквизиты" in value:
                    found = (ws, cell.row, cell.column)
                    break
            if found:
                break
        if found:
            break

    if not found:
        return False, "Не нашёл блок «Банковские реквизиты»."

    ws, start_row, start_col = found
    max_col = max(ws.max_column, start_col + 44)
    end_col = max_col
    mid_col = start_col + max(10, (end_col - start_col + 1) // 2)

    # Удаляем старый банковский блок строго от найденной строки до конца листа.
    rows_to_delete = ws.max_row - start_row + 1
    if rows_to_delete > 0:
        _unmerge_intersections(ws, start_row, ws.max_row, 1, ws.max_column)
        ws.delete_rows(start_row, rows_to_delete)

    # Берём базовый стиль рядом с местом вставки, чтобы блок не выглядел инородно.
    template_row = max(1, start_row - 1)
    base_cell = ws.cell(template_row, start_col)
    border = Border()
    align = Alignment(wrap_text=True, vertical="top", horizontal="left")
    font = copy(base_cell.font) if base_cell.font else Font(name="Calibri", size=10)
    if not font.sz:
        font = Font(name=font.name or "Calibri", size=10)

    row_count = max(len(PAYMENT_BANK_BLOCK_EN), len(PAYMENT_BANK_BLOCK_RU))
    for offset in range(row_count):
        r = start_row + offset
        left = ws.cell(r, start_col)
        right = ws.cell(r, mid_col)
        left.value = PAYMENT_BANK_BLOCK_EN[offset] if offset < len(PAYMENT_BANK_BLOCK_EN) else ""
        right.value = PAYMENT_BANK_BLOCK_RU[offset] if offset < len(PAYMENT_BANK_BLOCK_RU) else ""
        is_heading = offset in (0, len(PAYMENT_BANK_BLOCK_EN) - 2)
        for c in range(start_col, end_col + 1):
            cell = ws.cell(r, c)
            cell.alignment = align
            cell_font = copy(font)
            cell_font.bold = is_heading
            cell.font = cell_font
            cell.border = border
        ws.row_dimensions[r].height = 15

    # Две широкие колонки через merge: английский блок слева, русский справа.
    _unmerge_intersections(ws, start_row, start_row + row_count - 1, start_col, end_col)
    for r in range(start_row, start_row + row_count):
        ws.merge_cells(start_row=r, start_column=start_col, end_row=r, end_column=mid_col - 1)
        ws.merge_cells(start_row=r, start_column=mid_col, end_row=r, end_column=end_col)

    ws.column_dimensions[ws.cell(1, mid_col).column_letter].width = max(
        ws.column_dimensions[ws.cell(1, mid_col).column_letter].width or 8, 18
    )
    # Банковский блок начинается в B, поэтому финально фиксируем первые колонки после всех merge/width-операций.
    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 3
    ws.column_dimensions["C"].width = 3

    out_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)
    return True, f"Заменил банковские реквизиты на листе «{ws.title}», строка {start_row}."


# =========================
# Мастер "Письмо об утере инвойса"
# =========================
SIGNATURE_DB = STORAGE_DIR / "signatures.sqlite3"
SIGNATURE_DIR = STORAGE_DIR / "signatures"
LOST_INVOICE_TEMPLATE = Path(__file__).parent.parent / "assets" / "lost_invoice_letter_template.docx"
LOST_INVOICE_DOC_TYPE = "lost_invoice_letter"

_LOST_STEPS = [
    ("company",      LostInvoiceLetter.waiting_company.state,      "Введите название компании на английском, как должно быть в письме."),
    ("country",      LostInvoiceLetter.waiting_country.state,      "Введите страну компании."),
    ("letter_no",    LostInvoiceLetter.waiting_letter_no.state,    "Введите номер письма. Если подходит автономер — напишите <b>-</b>."),
    ("letter_date",  LostInvoiceLetter.waiting_letter_date.state,  "Введите дату письма. Если подходит сегодняшняя дата — напишите <b>-</b>."),
    ("awb_no",       LostInvoiceLetter.waiting_awb.state,          "Введите номер авианакладной."),
    ("places",       LostInvoiceLetter.waiting_places.state,       "Введите количество мест."),
    ("weight_kg",    LostInvoiceLetter.waiting_weight.state,       "Введите вес, кг."),
    ("amount",       LostInvoiceLetter.waiting_amount.state,       "Введите сумму и валюту, например: <b>12 345,67 USD</b>."),
    ("invoice_no",   LostInvoiceLetter.waiting_invoice_no.state,   "Введите номер инвойса."),
    ("invoice_date", LostInvoiceLetter.waiting_invoice_date.state, "Введите дату инвойса."),
]

_LOST_NEXT_STATE = {
    _LOST_STEPS[i][1]: _LOST_STEPS[i + 1][1]
    for i in range(len(_LOST_STEPS) - 1)
}
_LOST_FIELD_BY_STATE = {state: field for field, state, _ in _LOST_STEPS}
_LOST_PROMPT_BY_STATE = {state: prompt for _, state, prompt in _LOST_STEPS}


def _init_signature_db() -> None:
    SIGNATURE_DIR.mkdir(parents=True, exist_ok=True)
    if _signature_db_url():
        import psycopg
        with psycopg.connect(_signature_db_url()) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    CREATE TABLE IF NOT EXISTS signature_blocks (
                        id BIGSERIAL PRIMARY KEY,
                        company_key TEXT NOT NULL,
                        company_name TEXT NOT NULL,
                        country TEXT NOT NULL,
                        doc_type TEXT NOT NULL,
                        image_path TEXT,
                        image_ext TEXT,
                        image_data BYTEA,
                        created_at TEXT NOT NULL,
                        updated_at TEXT NOT NULL,
                        UNIQUE(company_key, doc_type)
                    )
                    """
                )
            conn.commit()
    else:
        SIGNATURE_DB.parent.mkdir(parents=True, exist_ok=True)
        with sqlite3.connect(SIGNATURE_DB) as conn:
            conn.execute(
                """
                CREATE TABLE IF NOT EXISTS signature_blocks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    company_key TEXT NOT NULL,
                    company_name TEXT NOT NULL,
                    country TEXT NOT NULL,
                    doc_type TEXT NOT NULL,
                    image_path TEXT,
                    image_ext TEXT,
                    image_data BLOB,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL,
                    UNIQUE(company_key, doc_type)
                )
                """
            )
            cols = {r[1] for r in conn.execute("PRAGMA table_info(signature_blocks)").fetchall()}
            if "image_ext" not in cols:
                conn.execute("ALTER TABLE signature_blocks ADD COLUMN image_ext TEXT")
            if "image_data" not in cols:
                conn.execute("ALTER TABLE signature_blocks ADD COLUMN image_data BLOB")


def _signature_db_url() -> str | None:
    return os.getenv("DATABASE_URL") or os.getenv("POSTGRES_URL")


def _normalize_key(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-zа-яё0-9]+", "_", value, flags=re.IGNORECASE)
    return value.strip("_") or "company"


def _safe_filename(value: str, fallback: str = "file") -> str:
    value = re.sub(r'[,<>:"/\\|?*\n\r\t]+', "_", value).strip(" ._")
    return value[:90] or fallback


def _default_letter_no() -> str:
    return datetime.now().strftime("%y%m%d") + "01"


def _default_letter_date() -> str:
    return datetime.now().strftime("%d.%m.%y")


def _materialize_signature(company: str, country: str, image_data: bytes | memoryview | None, image_ext: str | None, image_path: str | None) -> Path | None:
    if image_data:
        company_key = _normalize_key(f"{company}_{country}")
        ext = image_ext or ".png"
        if not ext.startswith("."):
            ext = "." + ext
        sig_dir = SIGNATURE_DIR / company_key
        sig_dir.mkdir(parents=True, exist_ok=True)
        out = sig_dir / f"{LOST_INVOICE_DOC_TYPE}{ext}"
        out.write_bytes(bytes(image_data))
        return out
    if image_path:
        path = Path(image_path)
        return path if path.exists() else None
    return None


def _signature_for(company: str, country: str, doc_type: str = LOST_INVOICE_DOC_TYPE) -> Path | None:
    _init_signature_db()
    company_key = _normalize_key(f"{company}_{country}")
    if _signature_db_url():
        import psycopg
        with psycopg.connect(_signature_db_url()) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT company_name, country, image_data, image_ext, image_path
                    FROM signature_blocks
                    WHERE company_key = %s AND doc_type = %s
                    """,
                    (company_key, doc_type),
                )
                row = cur.fetchone()
    else:
        with sqlite3.connect(SIGNATURE_DB) as conn:
            row = conn.execute(
                """
                SELECT company_name, country, image_data, image_ext, image_path
                FROM signature_blocks
                WHERE company_key = ? AND doc_type = ?
                """,
                (company_key, doc_type),
            ).fetchone()
    if not row:
        return None
    return _materialize_signature(row[0], row[1], row[2], row[3], row[4])


def _known_signature_companies(doc_type: str = LOST_INVOICE_DOC_TYPE) -> list[dict]:
    _init_signature_db()
    if _signature_db_url():
        import psycopg
        with psycopg.connect(_signature_db_url()) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT id, company_name, country
                    FROM signature_blocks
                    WHERE doc_type = %s
                    ORDER BY lower(company_name), lower(country)
                    """,
                    (doc_type,),
                )
                rows = cur.fetchall()
    else:
        with sqlite3.connect(SIGNATURE_DB) as conn:
            rows = conn.execute(
                """
                SELECT id, company_name, country
                FROM signature_blocks
                WHERE doc_type = ?
                ORDER BY lower(company_name), lower(country)
                """,
                (doc_type,),
            ).fetchall()
    return [{"id": r[0], "company_name": r[1], "country": r[2]} for r in rows]


def _signature_company_by_id(sig_id: int, doc_type: str = LOST_INVOICE_DOC_TYPE) -> dict | None:
    _init_signature_db()
    if _signature_db_url():
        import psycopg
        with psycopg.connect(_signature_db_url()) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    SELECT id, company_name, country
                    FROM signature_blocks
                    WHERE id = %s AND doc_type = %s
                    """,
                    (sig_id, doc_type),
                )
                row = cur.fetchone()
    else:
        with sqlite3.connect(SIGNATURE_DB) as conn:
            row = conn.execute(
                """
                SELECT id, company_name, country
                FROM signature_blocks
                WHERE id = ? AND doc_type = ?
                """,
                (sig_id, doc_type),
            ).fetchone()
    if not row:
        return None
    return {"id": row[0], "company_name": row[1], "country": row[2]}


def _save_signature(company: str, country: str, image_path: Path, doc_type: str = LOST_INVOICE_DOC_TYPE) -> None:
    _init_signature_db()
    company_key = _normalize_key(f"{company}_{country}")
    image_ext = image_path.suffix.lower() or ".png"
    image_data = image_path.read_bytes()
    now = datetime.now().isoformat(timespec="seconds")
    if _signature_db_url():
        import psycopg
        with psycopg.connect(_signature_db_url()) as conn:
            with conn.cursor() as cur:
                cur.execute(
                    """
                    INSERT INTO signature_blocks
                        (company_key, company_name, country, doc_type, image_path,
                         image_ext, image_data, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON CONFLICT(company_key, doc_type) DO UPDATE SET
                        company_name = excluded.company_name,
                        country      = excluded.country,
                        image_path   = excluded.image_path,
                        image_ext    = excluded.image_ext,
                        image_data   = excluded.image_data,
                        updated_at   = excluded.updated_at
                    """,
                    (company_key, company, country, doc_type, str(image_path),
                     image_ext, image_data, now, now),
                )
            conn.commit()
    else:
        with sqlite3.connect(SIGNATURE_DB) as conn:
            conn.execute(
                """
                INSERT INTO signature_blocks
                    (company_key, company_name, country, doc_type, image_path,
                     image_ext, image_data, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                ON CONFLICT(company_key, doc_type) DO UPDATE SET
                    company_name = excluded.company_name,
                    country      = excluded.country,
                    image_path   = excluded.image_path,
                    image_ext    = excluded.image_ext,
                    image_data   = excluded.image_data,
                    updated_at   = excluded.updated_at
                """,
                (company_key, company, country, doc_type, str(image_path),
                 image_ext, image_data, now, now),
            )


async def _ask_lost_step(message: Message, state_name: str) -> None:
    await message.answer(_LOST_PROMPT_BY_STATE[state_name], parse_mode="HTML")


def _replace_docx_paragraph(paragraph, replacements: dict[str, str], signature_path: Path | None) -> None:
    text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
    if not text:
        return

    signature_placeholder = "(блок печати и подписи)"
    if signature_placeholder in text:
        new_text = text.replace(signature_placeholder, "").strip()
        paragraph.clear()
        if new_text:
            paragraph.add_run(new_text)
            paragraph.add_run().add_break()
        if signature_path and signature_path.exists():
            from docx.shared import Cm
            from PIL import Image as _PILImage

            # Держим блок подписи на первой странице: Word сам пагинирует,
            # поэтому ограничиваем и ширину, и высоту картинки заранее.
            max_w_cm = 14.0
            max_h_cm = 7.2
            try:
                with _PILImage.open(str(signature_path)) as im:
                    w_px, h_px = im.size
                ratio = min(max_w_cm / w_px, max_h_cm / h_px)
                width_cm = max(1.0, w_px * ratio)
                height_cm = max(1.0, h_px * ratio)
            except Exception:
                width_cm = 10.0
                height_cm = None

            run = paragraph.add_run()
            if height_cm is None:
                run.add_picture(str(signature_path), width=Cm(width_cm))
            else:
                run.add_picture(str(signature_path), width=Cm(width_cm), height=Cm(height_cm))
        return

    new_text = text
    for src, dst in replacements.items():
        new_text = new_text.replace(src, dst)
    if new_text != text:
        paragraph.clear()
        paragraph.add_run(new_text)


def _fill_lost_invoice_template(data: dict, signature_path: Path, out_path: Path) -> None:
    from docx import Document

    doc = Document(str(LOST_INVOICE_TEMPLATE))
    company_country = f"{data['company']}, {data['country']}"
    replacements = {
        "(YYMMDD01)": data["letter_no"],
        "(DDMMYY)": data["letter_date"],
        "(Company name, Country)": company_country,
        "(Company name, country)": company_country,
        "(номер авианакладной)": data["awb_no"],
        "(кол-во мест)": data["places"],
        "(вес, кг)": data["weight_kg"],
        "(Сумма, валюта)": data["amount"],
        "(номер инвойса)": data["invoice_no"],
        "(дата инвойса)": data["invoice_date"],
    }

    for p in doc.paragraphs:
        _replace_docx_paragraph(p, replacements, signature_path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_docx_paragraph(p, replacements, signature_path)

    _remove_trailing_empty_docx_paragraphs(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _remove_trailing_empty_docx_paragraphs(doc) -> None:
    """Убирает пустые абзацы в конце шаблона, которые создают пустую страницу."""
    body = doc._body._element
    for child in list(body)[::-1]:
        if child.tag.endswith('}sectPr'):
            continue
        if not child.tag.endswith('}p'):
            break
        text = ''.join(child.itertext()).strip()
        has_drawing = bool(child.xpath('.//*[local-name()="drawing"]'))
        has_page_break = bool(child.xpath('.//*[local-name()="br" and @*[local-name()="type"]="page"]'))
        if text or has_drawing or has_page_break:
            break
        body.remove(child)


async def _finish_lost_invoice_letter(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    letter = data.get("lost_invoice", {})
    signature_path = _signature_for(letter["company"], letter["country"])
    if not signature_path:
        await state.set_state(LostInvoiceLetter.waiting_signature)
        await message.answer(
            "Для этой компании пока нет сохранённого блока печати и подписи.\n"
            "Пришлите картинку PNG/JPG — я сохраню её в базу и вставлю в письмо."
        )
        return

    user_dir = TMP_DIR / str(message.from_user.id) / "lost_invoice_letter"
    user_dir.mkdir(parents=True, exist_ok=True)
    out_name = (
        f"Письмо_об_утере_инвойса_"
        f"{_safe_filename(letter['company'])}_"
        f"{_safe_filename(letter['letter_date'])}.docx"
    )
    out_path = user_dir / out_name
    _fill_lost_invoice_template(letter, signature_path, out_path)
    await state.clear()
    await message.answer_document(
        FSInputFile(str(out_path)),
        caption="✅ Письмо об утере инвойса готово",
    )
    await message.answer(
        "Что делаем дальше?",
        reply_markup=get_staff_inline_menu(),
    )

# =========================
# Парсер спецификации
# =========================
# pypdf из Excel-PDF извлекает текст по-колоночно:
#   Блок A: строки "артикул  HS-код" — по порядку позиций
#   Блок B: строки данных "<№> <RU> / <EN> / <изготовитель> [qty] price amount"
# Стратегия: собрать articles[] из блока A, rows{} из блока B, соединить по индексу.

_ART_HS_RE  = re.compile(r'^([A-Z0-9][A-Za-z0-9\-]*RUO[C]?)\s+(\d{10})\s*$')
_ROW_START  = re.compile(r'^(\d+)\s+([А-яЁё]{3,})')
_PRICE_PAT  = r'\d{1,3}(?:\s\d{3})*,\d{2}'
_PRICE_END  = re.compile(rf'({_PRICE_PAT})\s+({_PRICE_PAT})\s*$')
_EN_KEYS    = ('antibod', 'monoclonal', 'polyclonal', 'concentrat', 'diluted',
               'detection', 'polymer', 'ready to use', 'anti-', ' kit', 'back-')

_RU_MONTHS = {
    "января": "January", "февраля": "February", "марта": "March",
    "апреля": "April",   "мая": "May",          "июня": "June",
    "июля": "July",      "августа": "August",   "сентября": "September",
    "октября": "October","ноября": "November",  "декабря": "December",
}

def _ru_date_to_en(date_ru: str) -> str:
    parts = date_ru.split()
    if len(parts) == 3:
        day, month_ru, year = parts
        month_en = _RU_MONTHS.get(month_ru.lower(), month_ru)
        return f"{month_en} {day}, {year}"
    return date_ru

def _ru_date_to_en_brief(date_ru: str) -> str:
    """'8 апреля 2026' → '8 April 2026' (day-first, for bilingual labels)."""
    parts = date_ru.split()
    if len(parts) == 3:
        day, month_ru, year = parts
        month_en = _RU_MONTHS.get(month_ru.lower(), month_ru)
        return f"{day} {month_en} {year}"
    return date_ru

def _has_cyrillic(s: str) -> bool:
    return bool(re.search('[А-яЁё]', s))

def _parse_price(s: str) -> float:
    return float(s.replace(' ', '').replace(',', '.'))

def _parse_last_line(last: str) -> tuple[float, float, int, str]:
    """
    Extract (price, amount, qty, manufacturer) from the final line of a row block.

    Core idea: find the two rightmost _PRICE_PAT tokens — they are (price_candidate, amount).
    Then resolve qty via:
      A) amount / price_candidate is a clean integer  → that's qty, price_candidate is price
      B) price_candidate has a leading digit group that is actually qty (e.g. "3 635,60"
         → qty=3, price=635,60). Try peeling it off and validate qty_b × price_b ≈ amount.
    """
    def _clean_mfr(s: str) -> str:
        s = re.sub(r'(\s+\d+)+$', '', s).strip()
        return '' if re.match(r'^\d+$', s) else s

    _PAT = re.compile(rf'({_PRICE_PAT})')
    matches = list(_PAT.finditer(last))
    if len(matches) < 2:
        return 0.0, 0.0, 1, last.strip()

    amount_m = matches[-1]
    price_m  = matches[-2]
    amount    = _parse_price(amount_m.group(1))
    price_raw = price_m.group(1)
    before    = last[:price_m.start()].strip()

    # Strategy A: price_raw is the actual price; qty = amount / price_raw
    price_a = _parse_price(price_raw)
    if price_a > 0:
        qty_float = amount / price_a
        qty_a = int(round(qty_float))
        if qty_a >= 1 and abs(qty_a * price_a - amount) < 0.02:
            return price_a, amount, qty_a, _clean_mfr(before)

    # Strategy B: price_raw has a leading digit that is actually qty
    # e.g. "3 635,60" → qty=3, price=635.60
    parts = price_raw.split()
    if len(parts) >= 2:
        try:
            qty_b  = int(parts[0])
            price_b = _parse_price(' '.join(parts[1:]))
            if price_b > 0 and qty_b >= 1 and abs(qty_b * price_b - amount) < 0.02:
                return price_b, amount, qty_b, _clean_mfr(before)
        except (ValueError, IndexError):
            pass

    return price_a, amount, 1, _clean_mfr(before)


def _split_row_block(first_rest: str, extra_lines: list[str]) -> dict:
    price = amount = 0.0
    qty = 1
    manufacturer = ''

    # Count trailing empty + lone-price lines — these are the "price tail"
    # (pypdf extracts qty/price/amount columns as separate lone tokens at the end of a row block)
    _LONE_RE = re.compile(rf'^({_PRICE_PAT})\s*$')
    n_price_lines = 0
    for el in reversed(extra_lines):
        if not el.strip() or _LONE_RE.match(el):
            n_price_lines += 1
        else:
            break

    if n_price_lines > 0:
        price_tail = extra_lines[len(extra_lines) - n_price_lines:]
        non_empty_tail = [el for el in price_tail if el.strip()]
        if len(non_empty_tail) >= 2:
            # Two lone-price tokens: treat as (price_candidate, amount)
            p, a, q, m = _parse_last_line(non_empty_tail[-2] + ' ' + non_empty_tail[-1])
            if a > 0:
                price, amount, qty, manufacturer = p, a, q, m
        elif non_empty_tail:
            p, a, q, m = _parse_last_line(non_empty_tail[-1])
            if a > 0:
                price, amount, qty, manufacturer = p, a, q, m

    # If price tail didn't yield amount: last extra_line may have both tokens
    # (this is the normal case when _PRICE_END terminated the block scan)
    if amount == 0.0 and extra_lines:
        p, a, q, m = _parse_last_line(extra_lines[-1])
        if a > 0:
            price, amount, qty, manufacturer = p, a, q, m
            n_price_lines = max(n_price_lines, 1)

    # If still no amount: entire row is on first_rest line (price suffix at end)
    if amount == 0.0:
        m_end = _PRICE_END.search(first_rest)
        if m_end:
            p, a, q, m = _parse_last_line(first_rest[m_end.start():])
            if a > 0:
                price, amount, qty, manufacturer = p, a, q, m
                first_rest = re.sub(r'\s+\d+\s*$', '', first_rest[:m_end.start()]).strip()

    middle_lines = extra_lines[:len(extra_lines) - n_price_lines] if extra_lines else []

    ru_parts: list[str] = []
    en_parts: list[str] = []

    if ' / ' in first_rest:
        slash = first_rest.index(' / ')
        ru_parts.append(first_rest[:slash].strip())
        en_after = first_rest[slash + 3:].strip()
        if en_after:
            en_parts.append(en_after)
    elif first_rest:
        ru_parts.append(first_rest)

    for line in middle_lines:
        if not line:
            continue
        if _has_cyrillic(line):
            ru_parts.append(line)
        else:
            en_parts.append(line)

    if not manufacturer and en_parts:
        full_en = ' '.join(en_parts)
        last_key_pos = -1
        last_key_len = 0
        for k in _EN_KEYS:
            pos = full_en.lower().rfind(k)
            if pos > last_key_pos:
                last_key_pos = pos
                last_key_len = len(k)
        if last_key_pos != -1:
            tail = re.sub(r'^[,;\s]+', '', full_en[last_key_pos + last_key_len:]).strip()
            _NOT_MFR = ('antibod', 'antibody', 'clone', 'concentration', 'volume')
            if tail and len(tail) <= 40 and not any(t in tail.lower() for t in _NOT_MFR):
                manufacturer = tail
                en_parts = [full_en[:last_key_pos + last_key_len].rstrip(' ,;')]

    return {
        'ru_desc':      ' '.join(ru_parts).strip(),
        'en_desc':      ' '.join(en_parts).strip(),
        'manufacturer': manufacturer,
        'qty':          qty,
        'price':        price,
        'amount':       amount,
    }


def _parse_spec_pdf(path: Path) -> dict | None:
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        full_text = "\n".join(p.extract_text() or "" for p in reader.pages)
    except Exception:
        return None

    lines = [ln.strip() for ln in full_text.split('\n')]

    # --- Header ---
    m_no       = re.search(r'Спецификация № (\d+)', full_text)
    m_date     = re.search(r'Спецификация № \d+ от (.+?) г\.', full_text)
    m_contract = re.search(r'Контракт № (.+?) /', full_text)
    m_idn      = re.search(r'ИДН: ([\d.]+)', full_text)
    m_buyer    = re.search(r'Покупатель / Buyer:\s*(.+)', full_text)
    m_seller   = re.search(r'Продавец / Seller:\s*(.+)', full_text)

    spec_no  = m_no.group(1)                                    if m_no       else '?'
    date_s   = m_date.group(1).strip()                         if m_date     else ''
    contract = m_contract.group(1).strip().lstrip('№').strip() if m_contract else ''
    idn      = m_idn.group(1).rstrip('.')                      if m_idn      else ''
    buyer    = m_buyer.group(1).strip()                        if m_buyer    else ''
    seller   = m_seller.group(1).strip()                       if m_seller   else ''

    # Single pass: collect articles (Block A) and row data blocks (Block B) from all lines.
    # Multi-page PDFs interleave article and row blocks per page, so we can't restrict
    # the row scan to "after the last article line".
    articles: list[tuple[str, str]] = []
    rows_dict: dict[int, dict] = {}
    row_line_idx: dict[int, int] = {}  # row_no -> index in lines[]

    i = 0
    while i < len(lines):
        line = lines[i]
        ma = _ART_HS_RE.match(line)
        if ma:
            articles.append((ma.group(1), ma.group(2)))
            i += 1
            continue
        ms = _ROW_START.match(line)
        if ms:
            row_no = int(ms.group(1))
            row_line_idx[row_no] = i
            first_rest = line[ms.start(2):].strip()
            block: list[str] = []
            j = i + 1
            while j < len(lines) and len(block) < 20:
                block.append(lines[j])
                if _PRICE_END.search(lines[j]):
                    break
                j += 1
            new_data = _split_row_block(first_rest, block)
            existing = rows_dict.get(row_no)
            # Prefer the entry with actual price data; false matches produce amount=0
            if not existing or (existing.get('amount', 0.0) == 0.0 and new_data.get('amount', 0.0) > 0.0):
                rows_dict[row_no] = new_data
            # Advance by 1 only — do NOT jump past the block.
            i += 1
            continue
        i += 1

    # Backward scan: pypdf sometimes extracts the "Amount" column BEFORE the row's
    # text block. For rows still missing amount, look for lone price tokens that
    # appear between the previous row's start and this row's start.
    _LONE = re.compile(rf'^({_PRICE_PAT})\s*$')
    for row_no, rd in list(rows_dict.items()):
        if rd.get('amount', 0.0) > 0:
            continue
        start = row_line_idx.get(row_no, -1)
        if start < 0:
            continue

        # Scan range: from (previous row's start + 1) to (this row's start - 1)
        prev_start = row_line_idx.get(row_no - 1, -1)
        back_limit = prev_start + 1 if prev_start >= 0 else max(0, start - 30)

        backward: list[float] = []
        for k in range(start - 1, back_limit - 1, -1):
            m = _LONE.match(lines[k])
            if m:
                backward.append(_parse_price(m.group(1)))

        # Also collect lone price tokens from the forward block
        next_start = row_line_idx.get(row_no + 1, len(lines))
        forward: list[float] = []
        for k in range(start + 1, min(next_start, start + 25)):
            m = _LONE.match(lines[k])
            if m:
                forward.append(_parse_price(m.group(1)))

        # Validate: find backward candidate that equals qty × forward price
        if forward:
            price = forward[0]
            if price > 0:
                for amt in backward:
                    qty_f = amt / price
                    qty = int(round(qty_f))
                    if qty >= 1 and abs(qty * price - amt) < 0.02:
                        rows_dict[row_no] = {**rd, 'amount': amt, 'price': price, 'qty': qty}
                        break
        elif len(backward) >= 2:
            # Both amount and price are backward (no forward price found)
            amt = backward[0]
            for price in backward[1:]:
                if price > 0:
                    qty_f = amt / price
                    qty = int(round(qty_f))
                    if qty >= 1 and abs(qty * price - amt) < 0.02:
                        rows_dict[row_no] = {**rd, 'amount': amt, 'price': price, 'qty': qty}
                        break

    if not articles:
        return None

    # --- Assemble rows: articles[idx] matched with rows_dict[idx+1] by index ---
    rows: list[dict] = []
    for idx, (article, hs_code) in enumerate(articles):
        rd = rows_dict.get(idx + 1, {})
        rows.append({
            'ru_desc':      rd.get('ru_desc', ''),
            'en_desc':      rd.get('en_desc', ''),
            'article':      article,
            'hs_code':      hs_code,
            'manufacturer': rd.get('manufacturer', ''),
            'qty':          rd.get('qty', 1),
            'price':        rd.get('price', 0.0),
            'amount':       rd.get('amount', 0.0),
        })

    if not rows:
        return None

    return {
        'spec_no':  spec_no,
        'date':     date_s,
        'contract': contract,
        'idn':      idn,
        'buyer':    buyer,
        'seller':   seller,
        'rows':     rows,
    }


def _find_col(ws, header_row: int, *keywords: str) -> int | None:
    """Return the first column in header_row whose cell text contains any keyword (case-insensitive)."""
    for c in range(1, ws.max_column + 1):
        v = str(ws.cell(row=header_row, column=c).value or '').lower()
        if any(kw.lower() in v for kw in keywords):
            return c
    return None


def _parse_spec_xlsx(path: Path) -> dict | None:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), data_only=True)
        ws = wb.active
    except Exception:
        return None

    # Locate header row: first row containing 'артикул' or 'article'
    header_row = None
    for r in range(1, min(40, ws.max_row + 1)):
        for c in range(1, ws.max_column + 1):
            v = str(ws.cell(row=r, column=c).value or '').lower()
            if 'артикул' in v or ('article' in v and 'item' not in v):
                header_row = r
                break
        if header_row:
            break
    if not header_row:
        return None

    # Map columns by keyword — resilient to column count/order changes within a section
    col_no    = _find_col(ws, header_row, '№', 'п/п', 'n°')
    col_ru    = _find_col(ws, header_row, 'товар', 'услуга', 'наименование')
    col_en    = _find_col(ws, header_row, 'description', 'описание')
    col_art   = _find_col(ws, header_row, 'артикул', 'article')
    col_hs    = _find_col(ws, header_row, 'тнвэд', 'hs code', 'hs-code', 'hscode')
    col_mfr   = _find_col(ws, header_row, 'изготовитель', 'manufactur')
    col_qty   = _find_col(ws, header_row, 'кол-во', 'кол.', 'quantity', 'qty')
    col_price = _find_col(ws, header_row, 'цена', 'price')
    col_amt   = _find_col(ws, header_row, 'сумма', 'amount')

    # Skip "column numbering" sub-row common in 1С exports.
    # Cells may be int, float, or string "1"/"2"/... — all must look like positive integers.
    sub_row = header_row + 1
    is_numbering = True
    has_any_val = False
    for c in range(1, ws.max_column + 1):
        v = ws.cell(row=sub_row, column=c).value
        if v is None:
            continue
        has_any_val = True
        try:
            f = float(v)
            if f != int(f) or f < 1:
                is_numbering = False
                break
        except (ValueError, TypeError):
            is_numbering = False
            break
    data_start = header_row + 2 if (has_any_val and is_numbering) else header_row + 1

    # Extract header metadata from pre-header rows
    pre_lines: list[str] = []
    for r in range(1, header_row):
        parts = []
        for c in range(1, ws.max_column + 1):
            v = ws.cell(row=r, column=c).value
            if v is not None:
                parts.append(str(v))
        if parts:
            pre_lines.append(' '.join(parts))
    full_pre = '\n'.join(pre_lines)

    m_no   = re.search(r'Спецификация\s*№\s*(\d+)', full_pre)
    m_date = re.search(r'Спецификация\s*№\s*\d+\s+от\s+(.+?)\s+г', full_pre)

    # Contract: «Контракт № X / Contract...» OR «к контракту № X»
    m_contract = (
        re.search(r'Контракт\s*[№#]+\s*[№]?\s*(.+?)\s*(?:Contract\b|\n|$)', full_pre, re.IGNORECASE) or
        re.search(r'контракт[ау]\s*[№#]?\s*(\S+)', full_pre, re.IGNORECASE)
    )
    m_idn = re.search(r'ИДН:\s*([\d.]+)', full_pre)

    # Buyer: «Покупатель / Buyer: VALUE» OR «Грузополучатель / Покупатель VALUE» (value in next cell)
    m_buyer = (
        re.search(r'Покупатель\s*/\s*Buyer\s*:?\s*(.+)',            full_pre) or
        re.search(r'Грузополучатель\s*/\s*Покупатель\s+(.+)',       full_pre) or
        re.search(r'Buyer\s*/\s*Покупатель\s*:?\s*(.+)',            full_pre, re.IGNORECASE)
    )
    # Seller: «Продавец / Seller: VALUE» OR «Грузоотправитель / Продавец VALUE»
    m_seller = (
        re.search(r'Продавец\s*/\s*Seller\s*:?\s*(.+)',             full_pre) or
        re.search(r'Грузоотправитель\s*/\s*Продавец\s+(.+)',        full_pre) or
        re.search(r'Seller\s*/\s*Продавец\s*:?\s*(.+)',             full_pre, re.IGNORECASE)
    )

    spec_no  = m_no.group(1)                                    if m_no       else '?'
    date_s   = m_date.group(1).strip()                         if m_date     else ''
    contract = m_contract.group(1).strip().lstrip('№').strip() if m_contract else ''
    idn      = m_idn.group(1).rstrip('.')                      if m_idn      else ''
    buyer    = m_buyer.group(1).strip()                        if m_buyer    else ''
    seller   = m_seller.group(1).strip()                       if m_seller   else ''

    def _cell_str(r: int, c: int | None) -> str:
        if c is None:
            return ''
        v = ws.cell(row=r, column=c).value
        return str(v).strip() if v is not None else ''

    def _cell_float(r: int, c: int | None) -> float:
        if c is None:
            return 0.0
        v = ws.cell(row=r, column=c).value
        if isinstance(v, (int, float)):
            return float(v)
        if isinstance(v, str):
            try:
                return float(v.replace('\xa0', '').replace(' ', '').replace(',', '.'))
            except ValueError:
                pass
        return 0.0

    # Keywords that mark non-data rows (subtotals, tax lines, totals)
    _SKIP_KW = ('итого', 'без налога', 'налог', 'total amount', 'total amt', 'nds', 'vat')

    def _row_text(r: int) -> str:
        """Concatenated lowercased text of all cells in a row."""
        return ' '.join(
            str(ws.cell(row=r, column=c).value or '').lower()
            for c in range(1, ws.max_column + 1)
        )

    rows: list[dict] = []
    doc_total: float | None = None

    for r in range(data_start, ws.max_row + 1):
        rt = _row_text(r)

        # Subtotal / tax / total rows — skip as data, but extract Total Amount if present
        if any(kw in rt for kw in _SKIP_KW):
            if 'total amount' in rt or 'total amt' in rt:
                # Numeric value for Total Amount is in col_amt or the first numeric cell in the row
                v = _cell_float(r, col_amt)
                if v <= 0:
                    for c in range(1, ws.max_column + 1):
                        raw = ws.cell(row=r, column=c).value
                        if isinstance(raw, (int, float)) and raw > 0:
                            v = float(raw)
                            break
                if v > 0:
                    doc_total = v
            continue

        if col_no:
            no_val = ws.cell(row=r, column=col_no).value
            if no_val is None:
                continue
            try:
                if int(no_val) < 1:
                    continue
            except (ValueError, TypeError):
                continue

        qty_raw = _cell_float(r, col_qty)
        rows.append({
            'ru_desc':      _cell_str(r, col_ru),
            'en_desc':      _cell_str(r, col_en),
            'article':      _cell_str(r, col_art),
            'hs_code':      _cell_str(r, col_hs),
            'manufacturer': _cell_str(r, col_mfr),
            'qty':          int(qty_raw) if qty_raw >= 1 else 1,
            'price':        _cell_float(r, col_price),
            'amount':       _cell_float(r, col_amt),
        })

    if not rows:
        return None

    # Extract payment/delivery terms block (same text in all XLSX files).
    # Collect rows from "Условия оплаты и поставки" through "Срок поставки"
    # regardless of whether the block is a single merged cell or separate rows.
    _TERMS_START = ('условия оплаты и поставки', 'payment and delivery terms')
    _TERMS_END   = ('срок поставки', 'delivery time')
    terms_text = ''
    _collecting = False
    _collected: list[str] = []

    for r in range(data_start, ws.max_row + 1):
        rt = _row_text(r)
        if not _collecting and any(kw in rt for kw in _TERMS_START):
            _collecting = True
        if not _collecting:
            continue
        vals = [
            str(ws.cell(row=r, column=c).value).strip()
            for c in range(1, ws.max_column + 1)
            if ws.cell(row=r, column=c).value is not None
            and str(ws.cell(row=r, column=c).value).strip()
        ]
        for v in vals:
            for line in v.split('\n'):
                line = line.strip()
                if line:
                    _collected.append(line)
        if any(kw in rt for kw in _TERMS_END):
            terms_text = '\n'.join(_collected)
            break

    if _collecting and not terms_text and _collected:
        terms_text = '\n'.join(_collected)

    return {
        'spec_no':    spec_no,
        'date':       date_s,
        'contract':   contract,
        'idn':        idn,
        'buyer':      buyer,
        'seller':     seller,
        'rows':       rows,
        'doc_total':  doc_total,
        'terms_text': terms_text,
    }


def _parse_spec_file(path: Path) -> dict | None:
    """Dispatch to XLSX or PDF parser based on file extension."""
    if path.suffix.lower() in ('.xlsx', '.xls'):
        return _parse_spec_xlsx(path)
    return _parse_spec_pdf(path)


# =========================
# 1) "📦 Скачать остатки"
# =========================
@router.callback_query(F.data == "menu:stock")
async def on_stock_button(callback: CallbackQuery):
    candidates = list(STORAGE_DIR.glob("last_report.*"))
    if not candidates:
        await callback.message.answer(
            "Пока нет загруженного отчёта. Попроси администратора загрузить файл 🙂"
        )
        await callback.answer()
        return
    await callback.message.answer_document(
        document=FSInputFile(path=candidates[0]),
        caption="Последний отчёт",
    )
    await callback.answer()


# =========================
# 2) "❓ Задать вопрос"
# =========================
@router.callback_query(F.data == "menu:question")
async def on_question_button(callback: CallbackQuery):
    await callback.message.edit_text(
        "Выберите отдел:",
        reply_markup=get_question_inline_menu(),
    )
    await callback.answer()


# =========================
# 3) "👩‍💼 Для сотрудников"
# =========================
@router.callback_query(F.data == "menu:staff")
async def on_staff_button(callback: CallbackQuery):
    user_id = callback.from_user.id
    allowed = is_staff(user_id)
    log_staff_event(user_id, callback.from_user.username, "menu:staff", allowed)
    if not allowed:
        await deny_staff_access(callback, action="menu:staff")
        return
    await callback.message.edit_text(
        "👩‍💼 Раздел для сотрудников.\n\nВыберите инструмент:",
        reply_markup=get_staff_inline_menu(),
    )
    await callback.answer()


# =========================
# 4) "⬅ Назад"
# =========================
@router.callback_query(F.data == "menu:back")
async def on_back_button(callback: CallbackQuery):
    await callback.message.edit_text(
        "Привет! 👋\n\nЭто Ixora Bot.\nВыберите действие:",
        reply_markup=get_main_inline_menu(),
    )
    await callback.answer()


# =========================
# 5a) Загрузка остатков (staff only)
# =========================
@router.callback_query(F.data == "staff:upload_stock")
async def on_upload_stock_start(callback: CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    allowed = is_staff(user_id)
    log_staff_event(user_id, callback.from_user.username, "staff:upload_stock", allowed)
    if not allowed:
        await deny_staff_access(callback, action="staff:upload_stock")
        return
    await state.set_state(UploadStock.waiting_file)
    await callback.message.edit_text(
        "📤 Загрузка отчёта остатков\n\n"
        "Пришлите файл Excel (.xlsx / .xls) или CSV.\n"
        "Для отмены напишите: ОТМЕНА"
    )
    await callback.answer()


@router.message(UploadStock.waiting_file, F.document)
async def on_upload_stock_file(message: Message, bot: Bot, state: FSMContext):
    doc = message.document
    fname = (doc.file_name or "")
    if not fname.lower().endswith((".xlsx", ".xls", ".csv")):
        await message.answer("Пожалуйста, пришлите файл Excel (.xlsx / .xls) или CSV.")
        return

    ext = Path(fname).suffix.lower()
    tg_file = await bot.get_file(doc.file_id)
    temp_path = STORAGE_DIR / f"temp{ext}"
    await bot.download_file(tg_file.file_path, destination=str(temp_path))

    for old in STORAGE_DIR.glob("last_report.*"):
        old.unlink(missing_ok=True)

    final_path = STORAGE_DIR / f"last_report{ext}"
    temp_path.rename(final_path)

    await state.clear()
    await message.answer(
        "✅ Отчёт об остатках обновлён!\n"
        "Теперь любой пользователь может скачать его через меню."
    )


@router.message(UploadStock.waiting_file, F.text.lower() == "отмена")
async def on_upload_stock_cancel(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Загрузка отменена.")


# =========================
# 5b) Корректировка инвойса на оплату (staff only)
# =========================
@router.callback_query(F.data == "staff:adjust_payment_invoice")
async def on_adjust_payment_invoice_start(callback: CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    allowed = is_staff(user_id)
    log_staff_event(user_id, callback.from_user.username, "staff:adjust_payment_invoice", allowed)
    if not allowed:
        await deny_staff_access(callback, action="staff:adjust_payment_invoice")
        return

    await state.set_state(AdjustPaymentInvoice.waiting_file)
    await callback.message.edit_text(
        "🏦 Корректировка инвойса на оплату\n\n"
        "Пришлите XLSX-инвойс. Я найду блок <b>«Банковские реквизиты»</b>, "
        "удалю его до конца листа и вставлю новый двуязычный блок.\n\n"
        "Для отмены напишите: ОТМЕНА",
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(AdjustPaymentInvoice.waiting_file, F.document)
async def on_adjust_payment_invoice_file(message: Message, state: FSMContext):
    progress = await message.answer("⏳ Обрабатываю инвойс...")
    try:
        doc = message.document
        fname = doc.file_name or ""
        if not fname.lower().endswith(".xlsx"):
            await progress.edit_text("Пришлите именно XLSX-файл инвойса.")
            return

        user_dir = TMP_DIR / str(message.from_user.id) / "adjust_payment_invoice" / str(int(time.time() * 1000))
        user_dir.mkdir(parents=True, exist_ok=True)

        src_path = user_dir / fname
        tg_file = await message.bot.get_file(doc.file_id)
        await message.bot.download_file(tg_file.file_path, destination=str(src_path))

        out_name = f"{src_path.stem}_исправлен.xlsx"
        out_path = user_dir / _safe_filename(out_name)

        ok, info = _adjust_payment_invoice_xlsx(src_path, out_path)
    except Exception as e:
        logger.exception("Payment invoice adjustment failed")
        await progress.edit_text(f"❌ Не удалось обработать файл: {e}")
        return

    if not ok:
        await progress.edit_text(f"❌ {info}")
        return

    await state.clear()
    await progress.edit_text("✅ Инвойс скорректирован, отправляю файл...")
    await message.answer_document(
        FSInputFile(str(out_path)),
        caption=f"✅ Инвойс скорректирован\n{info}",
    )
    await message.answer("Что делаем дальше?", reply_markup=get_staff_inline_menu())


@router.message(AdjustPaymentInvoice.waiting_file, F.text.lower() == "отмена")
async def on_adjust_payment_invoice_cancel(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Корректировка отменена.", reply_markup=get_staff_inline_menu())


# =========================
# 5c) Письмо об утере инвойса (staff only)
# =========================
@router.callback_query(F.data == "staff:lost_invoice_letter")
async def on_lost_invoice_start(callback: CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    allowed = is_staff(user_id)
    log_staff_event(user_id, callback.from_user.username, "staff:lost_invoice_letter", allowed)
    if not allowed:
        await deny_staff_access(callback, action="staff:lost_invoice_letter")
        return

    await state.set_state(LostInvoiceLetter.waiting_company)
    await state.update_data(lost_invoice={})
    companies = _known_signature_companies()
    if companies:
        await callback.message.edit_text(
            "📝 Письмо об утере инвойса\n\n"
            "Выберите компанию из сохранённых или введите новую.\n"
            "Так мы не плодим дубли вроде <i>SLEE medical</i> и <i>Slee Medical</i>.",
            parse_mode="HTML",
            reply_markup=_lost_company_kb(companies),
        )
    else:
        await callback.message.edit_text(
            "📝 Письмо об утере инвойса\n\n"
            "Я задам вопросы по бланку, потом соберу DOCX.\n"
            "Если блок печати и подписи для этой компании уже сохранён — возьму его из базы.\n\n"
            "Введите название компании на английском, как должно быть в письме.",
            parse_mode="HTML",
        )
    await callback.answer()


@router.callback_query(StateFilter(LostInvoiceLetter.waiting_company), F.data.startswith("lost_company:"))
async def on_lost_company_select(callback: CallbackQuery, state: FSMContext):
    value = callback.data.split(":", 1)[1]
    if value == "new":
        await callback.message.edit_text(
            "Введите название новой компании на английском, как должно быть в письме."
        )
        await callback.answer()
        return

    try:
        sig_id = int(value)
    except ValueError:
        await callback.answer("Не понял выбор компании", show_alert=True)
        return

    company = _signature_company_by_id(sig_id)
    if not company:
        await callback.answer("Компания не найдена, введите её вручную", show_alert=True)
        await callback.message.answer("Введите название компании на английском, как должно быть в письме.")
        return

    letter = {
        "company": company["company_name"],
        "country": company["country"],
    }
    await state.update_data(lost_invoice=letter)
    await state.set_state(LostInvoiceLetter.waiting_letter_no)
    await callback.message.edit_text(
        f"Компания: <b>{company['company_name']}, {company['country']}</b>\n\n"
        f"{_LOST_PROMPT_BY_STATE[LostInvoiceLetter.waiting_letter_no.state]}",
        parse_mode="HTML",
    )
    await callback.answer()


@router.message(
    StateFilter(
        LostInvoiceLetter.waiting_company,
        LostInvoiceLetter.waiting_country,
        LostInvoiceLetter.waiting_letter_no,
        LostInvoiceLetter.waiting_letter_date,
        LostInvoiceLetter.waiting_awb,
        LostInvoiceLetter.waiting_places,
        LostInvoiceLetter.waiting_weight,
        LostInvoiceLetter.waiting_amount,
        LostInvoiceLetter.waiting_invoice_no,
        LostInvoiceLetter.waiting_invoice_date,
    ),
    F.text,
)
async def on_lost_invoice_text(message: Message, state: FSMContext):
    if message.text and message.text.lower().strip() == "отмена":
        await state.clear()
        await message.answer("Ок, отменил создание письма.", reply_markup=get_staff_inline_menu())
        return

    current = await state.get_state()
    field = _LOST_FIELD_BY_STATE.get(current or "")
    if not field:
        await message.answer("Не понял, на каком шаге мы находимся. Начните мастер заново.")
        await state.clear()
        return

    value = (message.text or "").strip()
    if field == "letter_no" and value == "-":
        value = _default_letter_no()
    elif field == "letter_date" and value == "-":
        value = _default_letter_date()

    data = await state.get_data()
    letter = data.get("lost_invoice", {})
    letter[field] = value
    await state.update_data(lost_invoice=letter)

    next_state = _LOST_NEXT_STATE.get(current or "")
    if next_state:
        await state.set_state(next_state)
        await _ask_lost_step(message, next_state)
        return

    await _finish_lost_invoice_letter(message, state)


@router.message(LostInvoiceLetter.waiting_signature, F.photo | F.document)
async def on_lost_invoice_signature(message: Message, state: FSMContext):
    data = await state.get_data()
    letter = data.get("lost_invoice", {})
    if not letter:
        await state.clear()
        await message.answer("Данные письма потерялись. Запустите мастер заново.", reply_markup=get_staff_inline_menu())
        return

    company_key = _normalize_key(f"{letter['company']}_{letter['country']}")
    sig_dir = SIGNATURE_DIR / company_key
    sig_dir.mkdir(parents=True, exist_ok=True)

    if message.photo:
        file_id = message.photo[-1].file_id
        ext = ".jpg"
    else:
        doc = message.document
        fname = (doc.file_name or "").lower()
        ext = Path(fname).suffix.lower()
        if ext not in (".png", ".jpg", ".jpeg"):
            await message.answer("Нужна картинка PNG/JPG с блоком печати и подписи.")
            return
        file_id = doc.file_id

    tg_file = await message.bot.get_file(file_id)
    sig_path = sig_dir / f"{LOST_INVOICE_DOC_TYPE}{ext}"
    await message.bot.download_file(tg_file.file_path, destination=str(sig_path))
    _save_signature(letter["company"], letter["country"], sig_path)

    await message.answer("✅ Блок печати и подписи сохранил в базу. Собираю письмо...")
    await _finish_lost_invoice_letter(message, state)


@router.message(LostInvoiceLetter.waiting_signature)
async def on_lost_invoice_signature_wrong(message: Message):
    await message.answer("Пришлите картинку PNG/JPG с блоком печати и подписи.")


# =========================
# 5d) Запуск мастера
# =========================
@router.callback_query(F.data == "staff:merge_specs")
async def on_merge_specs_start(callback: CallbackQuery, state: FSMContext):
    user_id = callback.from_user.id
    allowed = is_staff(user_id)
    log_staff_event(user_id, callback.from_user.username, "staff:merge_specs", allowed)
    if not allowed:
        await deny_staff_access(callback, action="staff:merge_specs")
        return

    tmp_dir = TMP_DIR / str(user_id) / "specs"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)

    await state.set_state(MergeSpec.waiting_pdfs)
    await state.update_data(spec_files=[], spec_infos=[], pdf_files=[], status_msg_id=None)
    await callback.message.edit_text(
        "🔗 Объединить спецификации\n\n"
        "Загрузите <b>XLSX-файлы</b> спецификаций (по одному на каждую) — "
        "из них берётся табличная часть и шапка.\n\n"
        "Для страницы с условиями оплаты, подписями и печатями — "
        "загрузите также <b>один PDF</b> (любую спецификацию, "
        "в которой блок печатей и подписей <b>не разделён</b> между страницами).\n\n"
        "Дождитесь, пока система примет все файлы.",
        parse_mode="HTML",
    )
    await callback.answer()


# =========================
# 6) Приём файлов — PDF (для парсинга) или XLSX (диагностика структуры)
# =========================


# =========================
# 6a) XLSX upload — временный обработчик для диагностики структуры
# =========================
async def _debounced_status(message: Message, state: FSMContext, ts: float) -> None:
    """Ждёт 2 с после последнего файла, затем шлёт/редактирует одно статус-сообщение."""
    await asyncio.sleep(2)
    data = await state.get_data()
    if data.get('last_file_ts') != ts:
        return  # пришёл ещё файл — эта задача устарела
    spec_infos: list[str]     = data.get('spec_infos', [])
    pdf_files:  list[str]     = data.get('pdf_files', [])
    status_msg_id: int | None = data.get('status_msg_id')
    status = _build_upload_status(spec_infos, pdf_files)
    if status_msg_id:
        try:
            await message.bot.edit_message_text(
                chat_id=message.chat.id,
                message_id=status_msg_id,
                text=status,
                parse_mode="HTML",
                reply_markup=_merge_upload_kb(),
            )
            return
        except Exception:
            pass
    sent = await message.answer(status, parse_mode="HTML", reply_markup=_merge_upload_kb())
    await state.update_data(status_msg_id=sent.message_id)


@router.message(MergeSpec.waiting_pdfs, F.document)
async def on_xlsx_upload(message: Message, state: FSMContext):
    doc = message.document
    fname = (doc.file_name or "").lower()

    uid = message.from_user.id
    tmp_dir = TMP_DIR / str(uid) / "specs"
    tmp_dir.mkdir(parents=True, exist_ok=True)
    save_path = tmp_dir / doc.file_name
    tg_file = await message.bot.get_file(doc.file_id)
    await message.bot.download_file(tg_file.file_path, destination=str(save_path))

    data = await state.get_data()
    spec_files: list[str]     = data.get('spec_files', [])
    spec_infos: list[str]     = data.get('spec_infos', [])
    pdf_files:  list[str]     = data.get('pdf_files', [])
    status_msg_id: int | None = data.get('status_msg_id')

    if fname.endswith('.pdf'):
        if str(save_path) not in pdf_files:
            pdf_files.append(str(save_path))
        await state.update_data(pdf_files=pdf_files)

    elif fname.endswith(('.xlsx', '.xls')):
        result = _parse_spec_xlsx(save_path)
        if result is None:
            # Показываем структуру файла для диагностики
            try:
                from openpyxl import load_workbook as _lwb
                _wb = _lwb(save_path, read_only=True, data_only=True)
                _ws = _wb.active
                diag_lines = [f"❌ Не удалось распарсить: <b>{doc.file_name}</b>", "",
                              f"Листов: {len(_wb.sheetnames)} ({', '.join(_wb.sheetnames)})",
                              f"Размер: {_ws.max_row} строк × {_ws.max_column} столбцов", "",
                              "<b>Первые 20 строк:</b>"]
                for r in range(1, min(21, (_ws.max_row or 0) + 1)):
                    cells = []
                    for c in range(1, min(6, (_ws.max_column or 0) + 1)):
                        v = _ws.cell(r, c).value
                        if v is not None:
                            cells.append(f"[{str(v)[:30]}]")
                    if cells:
                        diag_lines.append(f"  стр{r}: {' '.join(cells)}")
                _wb.close()
            except Exception as _e:
                diag_lines = [f"❌ Не удалось распарсить: {doc.file_name} ({_e})"]
            await message.answer('\n'.join(diag_lines), parse_mode="HTML")
            return
        if str(save_path) not in spec_files:
            spec_files.append(str(save_path))
            rows = result['rows']
            computed  = sum(r['amount'] for r in rows)
            doc_total = result.get('doc_total')
            if doc_total is not None and abs(computed - doc_total) > 0.02:
                flag = f" ⚠️(расх. ${abs(computed - doc_total):,.2f})"
            else:
                flag = ""
            spec_infos.append(
                f"  • №{result['spec_no']}: {len(rows)} поз. / ${computed:,.2f}{flag}"
            )
        await state.update_data(spec_files=spec_files, spec_infos=spec_infos)

    else:
        await message.answer("Пожалуйста, пришлите PDF или XLSX файл.")
        return

    # Debounce: ждём 2 с после последнего файла, затем шлём одно статус-сообщение
    now = time.time()
    await state.update_data(last_file_ts=now)
    asyncio.create_task(_debounced_status(message, state, now))


# =========================
# 6b) "ДЕБАГ" — показываем сырой текст первого PDF
# =========================
@router.message(MergeSpec.waiting_pdfs, F.text.lower() == "дебаг")
async def on_merge_debug(message: Message, state: FSMContext):
    data = await state.get_data()
    spec_files: list[str] = data.get('spec_files', [])
    if not spec_files:
        await message.answer("Сначала загрузите хотя бы один PDF.")
        return
    try:
        from pypdf import PdfReader
        reader = PdfReader(spec_files[0])
        text = "\n".join(p.extract_text() or "" for p in reader.pages)
        snippet = text[:3000]
    except Exception as e:
        await message.answer(f"Ошибка чтения: {e}")
        return
    await message.answer(
        f"📄 Файл: {Path(spec_files[0]).name}\n"
        f"Страниц: {len(reader.pages)}\n\n"
        f"Сырой текст (первые 3000 символов):\n\n{snippet}"
    )


# =========================
# 6c) "ДЕБАГ2" — показываем результат парсера (articles + rows)
# =========================
@router.message(MergeSpec.waiting_pdfs, F.text.lower() == "дебаг2")
async def on_merge_debug2(message: Message, state: FSMContext):
    data = await state.get_data()
    spec_files: list[str] = data.get('spec_files', [])
    if not spec_files:
        await message.answer("Сначала загрузите хотя бы один PDF.")
        return

    path = Path(spec_files[0])
    result = _parse_spec_file(path)
    if result is None:
        await message.answer(
            f"❌ Парсер вернул None для {path.name}\n\n"
            "Для PDF: проверьте ДЕБАГ (articles[] может быть пустым).\n"
            "Для XLSX: убедитесь, что есть строка-заголовок с «Артикул»."
        )
        return

    rows = result['rows']
    no_amt = [i + 1 for i, r in enumerate(rows) if r['amount'] == 0.0]
    lines = [
        f"<b>Файл:</b> {path.name}",
        f"<b>№:</b> {result['spec_no']}  <b>Дата:</b> {result['date']}",
        f"<b>Строк:</b> {len(rows)}  Без суммы: {len(no_amt)}",
        (f"amount=0 для поз.: {no_amt}" if no_amt else "✅ Все позиции с суммой"),
        "",
    ]
    for i, r in enumerate(rows[:10]):
        flag = " ⚠️" if r['amount'] == 0.0 else ""
        lines.append(
            f"<b>#{i+1}</b>{flag} Art: {r['article']}\n"
            f"  RU: {r['ru_desc'][:50]}\n"
            f"  EN: {r['en_desc'][:50]}\n"
            f"  Mfr: {r['manufacturer'] or '—'}  qty={r['qty']}  "
            f"price={r['price']:.2f}  amt={r['amount']:.2f}"
        )
    if len(rows) > 10:
        lines.append(f"... ещё {len(rows) - 10} строк")

    await message.answer('\n'.join(lines), parse_mode="HTML")

    # Diagnostic: for PDF files, show raw lines around _ROW_START for rows with amount=0
    if no_amt and path.suffix.lower() == '.pdf':
        try:
            from pypdf import PdfReader
            raw_lines = [
                ln.strip()
                for ln in "\n".join(p.extract_text() or "" for p in PdfReader(str(path)).pages).split('\n')
            ]
        except Exception:
            return
        debug_msgs = []
        for target_no in no_amt[:3]:
            for idx, ln in enumerate(raw_lines):
                ms = _ROW_START.match(ln)
                if ms and int(ms.group(1)) == target_no:
                    before = raw_lines[max(0, idx - 10): idx]
                    after  = raw_lines[idx: idx + 40]
                    debug_msgs.append(
                        f"Строка #{target_no} (idx {idx}) — 10 до / 40 после:\n"
                        f"--- ДО ---\n" +
                        "\n".join(f"  [{idx - len(before) + j}] {repr(s)}" for j, s in enumerate(before)) +
                        f"\n--- ПОСЛЕ ---\n" +
                        "\n".join(f"  [{idx + j}] {repr(s)}" for j, s in enumerate(after))
                    )
                    break
        if debug_msgs:
            await message.answer("🔍 Диагностика (amount=0):\n\n" + "\n\n".join(debug_msgs))


# =========================
# 7) Читаем файлы — общий хелпер + обработчики (текст и кнопка)
# =========================
async def _do_merge_ready(message: Message, state: FSMContext):
    data = await state.get_data()
    spec_files: list[str] = data.get('spec_files', [])

    if not spec_files:
        await message.answer("Нет загруженных файлов. Загрузите хотя бы один XLSX.")
        return

    await message.answer("⏳ Читаю спецификации...")

    parsed: list[dict] = []
    errors: list[str] = []
    for f in spec_files:
        result = _parse_spec_file(Path(f))
        if result and result['rows']:
            parsed.append(result)
        else:
            errors.append(Path(f).name)

    if not parsed:
        await message.answer(
            "❌ Не удалось распарсить ни одну спецификацию.\n"
            "Убедитесь, что загружены XLSX из 1С УНФ со строкой-заголовком «Артикул»."
        )
        return

    if errors:
        await message.answer(f"⚠️ Не удалось распарсить: {', '.join(errors)}")

    parsed.sort(key=lambda x: int(x['spec_no']) if x['spec_no'].isdigit() else 0)
    merged_no = "_".join(p['spec_no'] for p in parsed)

    lines = ["<b>Найдено спецификаций:</b>"]
    total_qty = 0
    total_amount = 0.0
    for p in parsed:
        qty = sum(r['qty'] for r in p['rows'])
        amount = sum(r['amount'] for r in p['rows'])
        total_qty += qty
        total_amount += amount
        lines.append(f"  • №{p['spec_no']}: {len(p['rows'])} поз., {qty} шт., ${amount:,.2f}")
    lines.append(
        f"\n<b>Итого:</b> {sum(len(p['rows']) for p in parsed)} поз., "
        f"{total_qty} шт., ${total_amount:,.2f}"
    )
    lines.append(f"\n<b>Номер итогового документа:</b> {merged_no}")
    lines.append("\nЧтобы изменить номер — напишите его.")

    await state.update_data(parsed=parsed, merged_no=merged_no)
    await state.set_state(MergeSpec.confirm)
    await message.answer('\n'.join(lines), parse_mode="HTML", reply_markup=_merge_confirm_kb())


async def _check_pdf_and_proceed(message: Message, state: FSMContext) -> None:
    """Проверяет PDF перед чтением спецификаций."""
    data = await state.get_data()
    pdf_files: list[str] = data.get('pdf_files', [])

    if not pdf_files:
        await message.answer(
            "⚠️ PDF не загружен!\n"
            "Загрузите PDF-спецификацию с блоком подписей и печатей "
            "(блок не должен быть разделён между страницами)."
        )
        return

    if len(pdf_files) == 1:
        await state.update_data(terms_pdf=pdf_files[0])
        await _do_merge_ready(message, state)
        return

    # Несколько PDF — просим выбрать
    b = InlineKeyboardBuilder()
    for i, pdf_path in enumerate(pdf_files):
        b.add(InlineKeyboardButton(
            text=f"📄 {Path(pdf_path).name}",
            callback_data=f"merge:pdf:{i}",
        ))
    b.adjust(1)
    await message.answer(
        f"Загружено {len(pdf_files)} PDF-файла. "
        "Выберите, какой использовать для блока подписей и печатей:",
        reply_markup=b.as_markup(),
    )


@router.message(MergeSpec.waiting_pdfs, F.text.lower() == "готово")
async def on_merge_ready(message: Message, state: FSMContext):
    await _check_pdf_and_proceed(message, state)


@router.callback_query(MergeSpec.waiting_pdfs, F.data == "merge:read")
async def on_merge_read_cb(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await _check_pdf_and_proceed(callback.message, state)


@router.callback_query(MergeSpec.waiting_pdfs, F.data.startswith("merge:pdf:"))
async def on_pdf_select(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    idx = int(callback.data.split(":")[-1])
    data = await state.get_data()
    pdf_files: list[str] = data.get('pdf_files', [])
    if idx >= len(pdf_files):
        await callback.answer("Неверный выбор.", show_alert=True)
        return
    await state.update_data(terms_pdf=pdf_files[idx])
    await _do_merge_ready(callback.message, state)


@router.callback_query(MergeSpec.waiting_pdfs, F.data == "merge:more")
async def on_merge_more_cb(callback: CallbackQuery):
    await callback.answer("Загружайте следующий файл 📎")


# =========================
# 8) Генерация Excel + PDF — общий хелпер + обработчики
# =========================
async def _generate_merge(message: Message, state: FSMContext, user_id: int | None = None):
    uid = user_id if user_id is not None else message.from_user.id
    data = await state.get_data()
    parsed: list[dict] = data['parsed']
    merged_no: str = data['merged_no']
    spec_files: list[str] = data['spec_files']

    await message.answer("⏳ Генерирую документы...")

    out_dir = TMP_DIR / str(uid) / "out"
    out_dir.mkdir(parents=True, exist_ok=True)

    # Объединяем строки со сквозной нумерацией
    all_rows: list[dict] = []
    for spec in parsed:
        for row in spec['rows']:
            all_rows.append(dict(row))
    for i, row in enumerate(all_rows, start=1):
        row['spec_no'] = i

    total_qty    = sum(r['qty'] for r in all_rows)
    total_amount = sum(r['amount'] for r in all_rows)

    header   = parsed[0]
    date_ru  = header.get('date', '')
    date_en  = _ru_date_to_en(date_ru)
    contract = header.get('contract', '')
    idn      = header.get('idn', '')
    buyer    = header.get('buyer', '')
    seller   = header.get('seller', '')

    # File naming: English if Ixora is buyer, Russian if Ixora is seller
    # Date formatted as "25_Марта_2026" / "25_March_2026"
    _date_parts = date_ru.split()
    if len(_date_parts) == 3:
        _d, _m, _y = _date_parts
        _date_file_ru = f"{_d}_{_m.capitalize()}_{_y}"
        _month_en = _RU_MONTHS.get(_m.lower(), _m)
        _date_file_en = f"{_d}_{_month_en}_{_y}"
    else:
        _date_file_ru = _date_file_en = date_ru.replace(' ', '_')

    if 'ixora' in buyer.lower():
        _base_name = f"Specification_{merged_no}_dtd_{_date_file_en}"
    else:
        _base_name = f"Спецификация_{merged_no}_от_{_date_file_ru}"

    # ---- 1) Excel ----
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, Border, Side
    except Exception:
        await message.answer("Не установлен openpyxl. Установите: pip install openpyxl")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "Specification"

    date_bilingual = f"{_ru_date_to_en_brief(date_ru)}/{date_ru}" if date_ru else merged_no
    ws.append(["", "SpecificationNo/Спецификация №", merged_no])
    ws.append(["", "dated/от",                       date_bilingual])
    ws.append(["", "Contract No./Контракт №",         contract])
    ws.append(["", "IDN / ИДН",                       idn])
    ws.append(["", "Buyer/Покупатель",  buyer])
    _r = ws.max_row
    ws.merge_cells(f"C{_r}:I{_r}")
    ws.cell(_r, 3).alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[_r].height = 60

    ws.append(["", "Seller / Продавец", seller])
    _r = ws.max_row
    ws.merge_cells(f"C{_r}:I{_r}")
    ws.cell(_r, 3).alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[_r].height = 60

    ws.append([])

    col_headers = [
        "№", "Товары (RU)", "Description (EN)",
        "Артикул / Item", "Код ТНВЭД / HS Code", "Изготовитель / Manufacturer",
        "Кол-во / Qty", "Цена / Price", "Сумма / Amount",
    ]
    ws.append(col_headers)
    header_row = ws.max_row

    for r in all_rows:
        ws.append([
            r['spec_no'], r['ru_desc'], r['en_desc'],
            r['article'], r['hs_code'], r['manufacturer'],
            r['qty'], r['price'], r['amount'],
        ])

    ws.append([])
    ws.append(["", "", "", "", "", "", total_qty, "Total amount", total_amount])
    table_end_row = ws.max_row  # last row that gets borders

    # Payment / delivery terms (same text in all XLSX specs)
    terms_text = next((p.get('terms_text', '') for p in parsed if p.get('terms_text')), '')
    if terms_text:
        ws.append([])
        for line in terms_text.split('\n'):
            line = line.strip()
            if line:
                ws.append(["", line])

    # --- Alignment & font ---
    wrap = Alignment(wrap_text=True, vertical="top")
    for row in ws.iter_rows(min_row=header_row, max_row=ws.max_row):
        for cell in row:
            cell.alignment = wrap
    for cell in ws[header_row]:
        cell.font = Font(bold=True)

    # --- Cell borders on the data table ---
    _thin = Side(style='thin')
    _border = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
    for row in ws.iter_rows(min_row=header_row, max_row=table_end_row):
        for cell in row:
            cell.border = _border

    # --- Column widths ---
    for i, w in enumerate([5, 55, 45, 14, 14, 28, 7, 10, 12], start=1):
        ws.column_dimensions[chr(64 + i)].width = w

    # --- Print settings: landscape, fit all columns on one page ---
    ws.page_setup.orientation = 'landscape'
    ws.sheet_properties.pageSetUpPr.fitToPage = True  # enables fit-to-page mode
    ws.page_setup.fitToWidth  = 1   # fit to 1 page wide
    ws.page_setup.fitToHeight = 0   # unlimited pages tall

    xlsx_path = out_dir / f"{_base_name}.xlsx"
    wb.save(xlsx_path)
    await message.answer_document(
        FSInputFile(str(xlsx_path)),
        caption="✅ Excel спецификации",
    )

    # ---- 2) PDF ----
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
            Image as RLImage, KeepTogether,
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
        from reportlab.pdfgen import canvas as rl_canvas
    except Exception as e:
        await message.answer(
            f"Не установлена библиотека: {e}\n"
            "Установите: pip install reportlab"
        )
        return

    _fonts_dir = Path(__file__).parent.parent / "fonts"
    _font_regular = _fonts_dir / "DejaVuSans.ttf"
    _font_bold    = _fonts_dir / "DejaVuSans-Bold.ttf"
    try:
        pdfmetrics.registerFont(TTFont("Arial", str(_font_regular)))
        pdfmetrics.registerFont(TTFont("Arial-Bold", str(_font_bold)))
    except Exception as _fe:
        await message.answer(f"Не удалось зарегистрировать шрифт: {_fe}")
        return

    result_pdf = out_dir / f"{_base_name}.pdf"

    # NumberedCanvas: collects all pages on first pass, then replays them with
    # header (pages 2+) and footer (all pages) drawn after total page count is known.
    def _make_canvas_class(spec_merged_no, spec_date_ru, spec_date_en):
        class _NumberedCanvas(rl_canvas.Canvas):
            def __init__(self, *args, **kwargs):
                rl_canvas.Canvas.__init__(self, *args, **kwargs)
                self._saved_pages: list[dict] = []

            def showPage(self):
                self._saved_pages.append(dict(self.__dict__))
                self._startPage()

            def save(self):
                total = len(self._saved_pages)
                for pg, state in enumerate(self._saved_pages, start=1):
                    self.__dict__.update(state)
                    w, h = self._pagesize
                    self.saveState()
                    # Footer — all pages
                    self.setFont("Arial", 7)
                    self.drawCentredString(
                        w / 2, 5 * mm,
                        f"стр. {pg} из {total} страниц"
                    )
                    # Header — pages 2+
                    if pg > 1:
                        hdr = (
                            f"Спецификация № {spec_merged_no} "
                            f"от {spec_date_ru} г.  /  "
                            f"Specification № {spec_merged_no} "
                            f"dated {spec_date_en}"
                        )
                        self.setFont("Arial", 8)
                        self.drawString(14 * mm, h - 8 * mm, hdr)
                        self.setLineWidth(0.3)
                        self.line(14 * mm, h - 10 * mm, w - 14 * mm, h - 10 * mm)
                    self.restoreState()
                    rl_canvas.Canvas.showPage(self)
                rl_canvas.Canvas.save(self)

        return _NumberedCanvas

    def _ps(name, size=9, leading=11, space_after=2, bold=False):
        return ParagraphStyle(
            name=name,
            parent=getSampleStyleSheet()["Normal"],
            fontName="Arial-Bold" if bold else "Arial",
            fontSize=size, leading=leading, spaceAfter=space_after,
        )

    title_s = _ps("T",  size=12, leading=15, space_after=1, bold=True)
    sub_s   = _ps("S",  size=11, leading=14, space_after=2)
    meta_s  = _ps("M",  size=9,  leading=11, space_after=2)
    cell_s  = _ps("C",  size=7,  leading=8.5, space_after=0)
    cell_sb = _ps("CB", size=7,  leading=8.5, space_after=0, bold=True)

    doc = SimpleDocTemplate(
        str(result_pdf),
        pagesize=landscape(A4),
        leftMargin=14*mm, rightMargin=14*mm,
        topMargin=16*mm,  bottomMargin=16*mm,
    )

    # --- Render terms section from PDF to image (before building story) ---
    # Portrait source → scale to fill landscape available width (~1.41x).
    # Landscape source → keep as-is (already landscape-sized).
    # KeepTogether ensures spacer+image land on the same page: last table page
    # if height fits there, otherwise a new page — NumberedCanvas covers both.
    terms_pdf_path = data.get('terms_pdf')
    terms_img_info = None  # (png_path, display_w_pt, display_h_pt)

    # Explicit layout constants — do NOT use doc.width/doc.height (those are
    # full page dimensions, not available frame dimensions).
    _LAND_W, _LAND_H = landscape(A4)
    _AVAIL_W = _LAND_W - 2 * 14 * mm   # landscape available width  (~762 pt)
    _AVAIL_H = _LAND_H - 2 * 16 * mm   # landscape available height (~504 pt)
    _SPACER_H = 8 * mm
    _FRAME_H  = _AVAIL_H - 12 - _SPACER_H   # minus Frame padding and spacer (~469 pt)

    if terms_pdf_path:
        try:
            import fitz as _fitz
            fdoc = _fitz.open(str(terms_pdf_path))
            t_idx = None
            for i in range(len(fdoc)):
                if any(n in fdoc[i].get_text().lower()
                       for n in ["условия оплаты и поставки", "payment and delivery terms"]):
                    t_idx = i
                    break
            if t_idx is None:
                for i in range(len(fdoc) - 1, -1, -1):
                    if fdoc[i].get_text().strip():
                        t_idx = i
                        break
            if t_idx is not None:
                page_f = fdoc[t_idx]
                pr = page_f.rect
                is_portrait = pr.height > pr.width

                # Text search to find exact crop start; tight fallback if not found
                clip_y0 = pr.height * 0.65  # fallback: bottom 35%
                for s in ["Условия оплаты и поставки", "Условия оплаты",
                          "Payment and delivery terms"]:
                    hits = page_f.search_for(s)
                    if hits:
                        clip_y0 = max(0.0, sorted(hits, key=lambda _r: _r.y0)[0].y0 - 5)
                        break

                # Find actual content bottom: last text block below clip_y0 + 60pt padding
                # to capture stamps and signature lines that sit below the last text row.
                text_blocks = [
                    b for b in page_f.get_text("blocks")
                    if b[1] >= clip_y0 - 5 and b[4].strip()
                ]
                if text_blocks:
                    clip_y1 = min(max(b[3] for b in text_blocks) + 60, pr.height)
                else:
                    clip_y1 = pr.height

                clip = _fitz.Rect(0, clip_y0, pr.width, clip_y1)
                pix = page_f.get_pixmap(
                    matrix=_fitz.Matrix(3.0, 3.0), clip=clip, alpha=False
                )
                fdoc.close()
                terms_png = out_dir / "terms_crop.png"
                pix.save(str(terms_png))

                # Portrait → 1.4× scale; landscape → keep native size; cap to frame
                if is_portrait:
                    disp_w = clip.width * 1.4
                    disp_h = clip.height * 1.4
                else:
                    disp_w = clip.width
                    disp_h = clip.height

                if disp_w > _AVAIL_W:       # cap width, preserve ratio
                    disp_h *= _AVAIL_W / disp_w
                    disp_w  = _AVAIL_W
                if disp_h > _FRAME_H:       # cap height, preserve ratio
                    disp_w *= _FRAME_H / disp_h
                    disp_h  = _FRAME_H

                terms_img_info = (terms_png, disp_w, disp_h)
            else:
                fdoc.close()
        except Exception as ex:
            await message.answer(f"⚠️ Не удалось подготовить страницу с условиями: {ex}")

    story = [
        Paragraph(f"Спецификация № {merged_no} от {date_ru} г. /", title_s),
        Paragraph(f"Specification № {merged_no} dated {date_en}", sub_s),
        Paragraph(
            f"Контракт № {contract} / Contract № {contract}", meta_s
        ),
        Paragraph(f"ИДН: {idn}.", meta_s),
        Paragraph(f"Покупатель / Buyer: {buyer}", meta_s),
        Paragraph(f"Продавец / Seller: {seller}", meta_s),
        Spacer(1, 4*mm),
    ]

    col_names = [
        "№",
        "Товар (Услуга)",
        "description",
        "Артикул/\nItem",
        "Код ТНВЭД/\nHS Code",
        "Изготовитель/\nManufacturer",
        "Кол-во/\nqty",
        "Цена/\nPrice",
        "Сумма/\nAmount",
    ]
    table_data = [[Paragraph(h, cell_sb) for h in col_names]]

    for r in all_rows:
        table_data.append([
            Paragraph(str(r['spec_no']), cell_s),
            Paragraph(r['ru_desc'],      cell_s),
            Paragraph(r['en_desc'],      cell_s),
            Paragraph(r['article'],      cell_s),
            Paragraph(r['hs_code'],      cell_s),
            Paragraph(r['manufacturer'], cell_s),
            Paragraph(str(r['qty']),     cell_s),
            Paragraph(f"{r['price']:.2f}",  cell_s),
            Paragraph(f"{r['amount']:.2f}", cell_s),
        ])

    # сумма долей = 1.00
    col_fracs = [0.03, 0.21, 0.19, 0.09, 0.08, 0.15, 0.05, 0.08, 0.09]
    col_w = [doc.width * f for f in col_fracs]

    tbl = Table(table_data, colWidths=col_w, repeatRows=1, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("ALIGN",        (0, 0), (0,  -1), "CENTER"),
        ("ALIGN",        (6, 0), (8,  -1), "RIGHT"),
        ("GRID",         (0, 0), (-1, -1), 0.3, colors.grey),
        ("LEFTPADDING",  (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING",   (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 2),
        ("BACKGROUND",   (0, 0), (-1,  0), colors.whitesmoke),
    ]))

    story.append(tbl)
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        f"Total amount: {total_amount:,.2f} USD",
        _ps("Tot", size=9, leading=11, bold=True),
    ))

    # KeepTogether: if spacer+image fit in remaining space on the last table page,
    # ReportLab places them there; otherwise they move to a new page together.
    # NumberedCanvas draws the correct header/footer on whichever page they land on.
    if terms_img_info:
        t_path, t_w, t_h = terms_img_info
        story.append(KeepTogether([
            Spacer(1, _SPACER_H),
            RLImage(str(t_path), width=t_w, height=t_h),
        ]))

    doc.build(
        story,
        canvasmaker=_make_canvas_class(merged_no, date_ru, date_en),
    )

    caption = (
        "✅ Итоговый PDF: таблица + условия и подписи"
        if terms_img_info else
        "✅ Итоговый PDF: только таблица (PDF с условиями не был загружен)"
    )
    await message.answer_document(
        FSInputFile(str(result_pdf)),
        caption=caption,
    )

    # --- Проверка: входящие спецификации vs итоговая ---
    in_pos = sum(len(p['rows']) for p in parsed)
    in_qty = sum(r['qty']    for p in parsed for r in p['rows'])
    in_amt = sum(r['amount'] for p in parsed for r in p['rows'])

    chk_lines = ["📊 <b>Проверка объединения:</b>\n"]
    chk_lines.append("<b>Входящие спецификации:</b>")
    for p in parsed:
        p_pos = len(p['rows'])
        p_qty = sum(r['qty']    for r in p['rows'])
        p_amt = sum(r['amount'] for r in p['rows'])
        chk_lines.append(f"  • №{p['spec_no']}: {p_pos} поз / {p_qty} шт / ${p_amt:,.2f}")
    chk_lines.append(
        f"\n<b>Итого вход:</b>  {in_pos} поз / {in_qty} шт / ${in_amt:,.2f}"
    )
    chk_lines.append(
        f"<b>Итоговая спец:</b> {len(all_rows)} поз / {total_qty} шт / ${total_amount:,.2f}"
    )
    all_match = (
        in_pos == len(all_rows)
        and in_qty == total_qty
        and abs(in_amt - total_amount) < 0.01
    )
    chk_lines.append("\n✅ Всё сошлось." if all_match else "\n⚠️ РАСХОЖДЕНИЕ! Проверьте данные.")

    await state.set_state(None)          # выходим из confirm — предотвращает петлю
    await state.update_data(last_out_dir=str(out_dir))

    await message.answer(
        '\n'.join(chk_lines) + "\n\nЕсли всё ок — нажмите <b>АРХИВ</b>.",
        parse_mode="HTML",
        reply_markup=_merge_done_kb(),
    )


# =========================
# 9) Запуск генерации PDF (текст или кнопка)
# =========================
@router.message(MergeSpec.confirm, F.text.lower().startswith("делаем"))
async def on_merge_generate(message: Message, state: FSMContext):
    await _generate_merge(message, state)


@router.callback_query(MergeSpec.confirm, F.data == "merge:generate")
async def on_merge_generate_cb(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await _generate_merge(callback.message, state, user_id=callback.from_user.id)


# =========================
# 10) Правим номер итогового документа (catch-all в confirm state)
# =========================
@router.message(MergeSpec.confirm)
async def on_merge_edit_number(message: Message, state: FSMContext):
    new_no = (message.text or "").strip()
    if not new_no:
        return
    await state.update_data(merged_no=new_no)
    await message.answer(
        f"✅ Номер изменён: <b>{new_no}</b>\nНапишите <b>ДЕЛАЕМ PDF</b> для генерации.",
        parse_mode="HTML",
    )


# =========================
# АРХИВ — переносим папку последнего запуска
# =========================
async def _do_archive(message: Message, state: FSMContext, user_id: int) -> None:
    if not is_staff(user_id):
        await deny_staff_access(message)
        return

    data = await state.get_data()
    last_out_dir = data.get("last_out_dir")
    if not last_out_dir:
        await message.answer("Не вижу, что архивировать 🙂 Сначала сделайте: ДЕЛАЕМ PDF")
        return

    src = Path(last_out_dir)
    if not src.exists():
        await message.answer("Папка результата уже не найдена.")
        return

    archive_root = STORAGE_DIR / "archive" / str(user_id)
    archive_root.mkdir(parents=True, exist_ok=True)
    existing = [p for p in archive_root.iterdir() if p.is_dir() and p.name.isdigit()]
    next_no = max((int(p.name) for p in existing), default=0) + 1
    dst = archive_root / str(next_no)

    try:
        shutil.move(str(src), str(dst))
    except Exception as e:
        await message.answer(f"Не смог переместить в архив: {type(e).__name__}: {e}")
        return

    await message.answer(
        f"👍 Перенёс в архив: `{dst}`",
        reply_markup=_archive_done_kb(),
    )
    await state.clear()


@router.message(F.text.lower().strip() == "архив")
async def on_archive(message: Message, state: FSMContext):
    await _do_archive(message, state, message.from_user.id)


@router.callback_query(F.data == "merge:archive")
async def on_archive_cb(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await _do_archive(callback.message, state, callback.from_user.id)


@router.callback_query(F.data == "merge:invoice")
async def on_invoice_cb(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await _do_invoice(callback.message, state)


async def _do_invoice(message: Message, state: FSMContext) -> None:
    data = await state.get_data()
    parsed: list[dict] = data.get("parsed", [])
    merged_no: str      = data.get("merged_no", "")
    last_out_dir        = data.get("last_out_dir")

    if not parsed or not merged_no or not last_out_dir:
        await message.answer("❌ Нет данных для генерации инвойса. Сначала сформируйте спецификацию.")
        return

    out_dir = Path(last_out_dir)

    all_rows: list[dict] = []
    for spec in parsed:
        for row in spec['rows']:
            all_rows.append(dict(row))
    for i, row in enumerate(all_rows, start=1):
        row['spec_no'] = str(i)

    total_qty    = sum(r['qty']    for r in all_rows)
    total_amount = sum(r['amount'] for r in all_rows)

    header   = parsed[0]
    date_ru  = header.get('date', '')
    date_en  = _ru_date_to_en(date_ru)
    buyer    = header.get('buyer',    '')
    seller   = header.get('seller',   '')
    contract = header.get('contract', '')
    idn      = header.get('idn',      '')

    _date_parts  = date_ru.split()
    _date_file_ru = "_".join(_date_parts) if _date_parts else merged_no
    _date_file_en = _ru_date_to_en_brief(date_ru).replace(" ", "_") if date_ru else merged_no

    _inv_base = (
        f"Invoice_{merged_no}_dtd_{_date_file_en}"
        if date_en else
        f"Инвойс_{merged_no}_от_{_date_file_ru}"
    )

    # ---- 1) Excel инвойс ----
    await message.answer("⏳ Генерирую инвойс...")
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, Border, Side
    except Exception:
        await message.answer("❌ openpyxl не установлен.")
        return

    wb = Workbook()
    ws = wb.active
    ws.title = "Invoice"

    date_bilingual = f"{_ru_date_to_en_brief(date_ru)}/{date_ru}" if date_ru else merged_no
    ws.append(["", "InvoiceNo/Инвойс №",           merged_no])
    ws.append(["", "dated/от",                      date_bilingual])
    ws.append(["", "Contract No./Контракт №",        contract])
    ws.append(["", "IDN / ИДН",                      idn])
    ws.append(["", "Buyer/Покупатель",  buyer])
    _r = ws.max_row
    ws.merge_cells(f"C{_r}:I{_r}")
    ws.cell(_r, 3).alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[_r].height = 60

    ws.append(["", "Seller / Продавец", seller])
    _r = ws.max_row
    ws.merge_cells(f"C{_r}:I{_r}")
    ws.cell(_r, 3).alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[_r].height = 60

    ws.append([])

    col_headers = [
        "№", "Товары (RU)", "Description (EN)",
        "Артикул / Item", "Код ТНВЭД / HS Code", "Изготовитель / Manufacturer",
        "Кол-во / Qty", "Цена / Price", "Сумма / Amount",
    ]
    ws.append(col_headers)
    header_row = ws.max_row

    for r in all_rows:
        ws.append([
            r['spec_no'], r['ru_desc'], r['en_desc'],
            r['article'], r['hs_code'], r['manufacturer'],
            r['qty'], r['price'], r['amount'],
        ])

    ws.append([])
    ws.append(["", "", "", "", "", "", total_qty, "Total amount", total_amount])
    table_end_row = ws.max_row

    terms_text = next((p.get('terms_text', '') for p in parsed if p.get('terms_text')), '')
    if terms_text:
        ws.append([])
        for line in terms_text.split('\n'):
            line = line.strip()
            if line:
                ws.append(["", line])

    wrap = Alignment(wrap_text=True, vertical="top")
    for row in ws.iter_rows(min_row=header_row, max_row=ws.max_row):
        for cell in row:
            cell.alignment = wrap
    for cell in ws[header_row]:
        cell.font = Font(bold=True)

    _thin = Side(style='thin')
    _border = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)
    for row in ws.iter_rows(min_row=header_row, max_row=table_end_row):
        for cell in row:
            cell.border = _border

    for i, w in enumerate([5, 55, 45, 14, 14, 28, 7, 10, 12], start=1):
        ws.column_dimensions[chr(64 + i)].width = w

    ws.page_setup.orientation = 'landscape'
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0

    inv_xlsx = out_dir / f"{_inv_base}.xlsx"
    wb.save(inv_xlsx)
    await message.answer_document(FSInputFile(str(inv_xlsx)), caption="✅ Excel инвойс")

    # ---- 2) PDF инвойс ----
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import (
            SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
            KeepTogether,
        )
        from reportlab.platypus import Image as RLImage
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
    except Exception as _e:
        await message.answer(f"❌ ReportLab не установлен: {_e}")
        return

    _fonts_dir   = Path(__file__).parent.parent / "fonts"
    _font_regular = _fonts_dir / "DejaVuSans.ttf"
    _font_bold    = _fonts_dir / "DejaVuSans-Bold.ttf"
    try:
        pdfmetrics.registerFont(TTFont("Arial", str(_font_regular)))
        pdfmetrics.registerFont(TTFont("Arial-Bold", str(_font_bold)))
    except Exception as _fe:
        await message.answer(f"❌ Шрифт: {_fe}")
        return

    inv_pdf = out_dir / f"{_inv_base}.pdf"

    _AVAIL_W = landscape(A4)[0] - 20 * mm
    _AVAIL_H = landscape(A4)[1] - 20 * mm
    _SPACER_H = 8 * mm
    _FRAME_H  = _AVAIL_H - 12 - _SPACER_H

    def _ps_inv(name, size=8, leading=10, bold=False, space_after=0):
        from reportlab.lib.styles import ParagraphStyle
        return ParagraphStyle(
            name,
            parent=getSampleStyleSheet()["Normal"],
            fontName="Arial-Bold" if bold else "Arial",
            fontSize=size, leading=leading, spaceAfter=space_after,
        )

    # Canvas с заголовком на каждой странице (кроме первой)
    def _make_inv_canvas(inv_no, d_ru, d_en):
        from reportlab.pdfgen.canvas import Canvas
        class _InvCanvas(Canvas):
            def __init__(self, filename, **kw):
                super().__init__(filename, **kw)
                self._saved_page_states = []
            def showPage(self):
                self._saved_page_states.append(dict(self.__dict__))
                self._startPage()
            def save(self):
                total = len(self._saved_page_states)
                for i, state in enumerate(self._saved_page_states, 1):
                    self.__dict__.update(state)
                    self._draw_inv_decorations(i, total, inv_no, d_ru, d_en)
                    super().showPage()
                super().save()
            def _draw_inv_decorations(self, page, total, inv_no, d_ru, d_en):
                w, h = landscape(A4)
                self.saveState()
                self.setFont("Arial", 7)
                self.drawCentredString(w / 2, 6 * mm, f"стр. {page} / {total}")
                if page > 1:
                    self.setFont("Arial", 8)
                    hdr = f"Инвойс / Invoice № {inv_no}  от / as of {d_ru} г. / {d_en}"
                    self.drawString(14 * mm, h - 8 * mm, hdr)
                self.restoreState()
        return _InvCanvas

    doc = SimpleDocTemplate(
        str(inv_pdf),
        pagesize=landscape(A4),
        leftMargin=10 * mm, rightMargin=10 * mm,
        topMargin=10 * mm, bottomMargin=10 * mm,
    )

    story = []

    # Заголовок
    title_text = (
        f"Инвойс / Invoice № {merged_no}  "
        f"от {date_ru} г. / as of {date_en}"
        if date_ru else
        f"Инвойс / Invoice № {merged_no}"
    )
    story.append(Paragraph(title_text, _ps_inv("Title", size=11, leading=14, bold=True, space_after=4)))

    # Метаданные
    meta_lines = []
    if contract: meta_lines.append(f"Контракт / Contract: {contract}")
    if idn:      meta_lines.append(f"ИДН / IDN: {idn}")
    if buyer:    meta_lines.append(f"Покупатель / Buyer: {buyer}")
    if seller:   meta_lines.append(f"Продавец / Seller: {seller}")
    for ln in meta_lines:
        story.append(Paragraph(ln, _ps_inv("Meta", size=8, leading=10, space_after=1)))
    story.append(Spacer(1, 3))

    # Таблица товаров
    col_hdrs = ["№", "Наименование / Description", "Description (EN)",
                 "Артикул\nItem", "ТНВЭД\nHS Code", "Изготовитель\nManufacturer",
                 "Кол-во\nQty", "Цена\nPrice", "Сумма\nAmount"]
    table_data = [col_hdrs]
    for r in all_rows:
        table_data.append([
            r['spec_no'],
            Paragraph(str(r['ru_desc'] or ''), _ps_inv("C")),
            Paragraph(str(r['en_desc'] or ''), _ps_inv("C")),
            Paragraph(str(r['article'] or ''), _ps_inv("C")),
            str(r['hs_code'] or ''),
            Paragraph(str(r['manufacturer'] or ''), _ps_inv("C")),
            r['qty'],
            f"{r['price']:,.2f}" if r['price'] else '',
            f"{r['amount']:,.2f}",
        ])

    col_fracs = [0.03, 0.21, 0.19, 0.09, 0.08, 0.15, 0.05, 0.08, 0.09]
    col_w = [doc.width * f for f in col_fracs]
    tbl = Table(table_data, colWidths=col_w, repeatRows=1, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("FONTNAME",     (0, 0), (-1, -1), "Arial"),
        ("FONTSIZE",     (0, 0), (-1, -1), 8),
        ("FONTNAME",     (0, 0), (-1,  0), "Arial-Bold"),
        ("VALIGN",       (0, 0), (-1, -1), "TOP"),
        ("ALIGN",        (0, 0), (0,  -1), "CENTER"),
        ("ALIGN",        (6, 0), (8,  -1), "RIGHT"),
        ("GRID",         (0, 0), (-1, -1), 0.3, colors.grey),
        ("LEFTPADDING",  (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING",   (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING",(0, 0), (-1, -1), 2),
        ("BACKGROUND",   (0, 0), (-1,  0), colors.whitesmoke),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        f"Total amount: {total_amount:,.2f} USD",
        _ps_inv("Tot", size=9, leading=11, bold=True),
    ))

    # Блок подписи: статичный PNG с реальной подписью, печатью и реквизитами
    _sig_png = Path(__file__).parent.parent / "assets" / "invoice_signature.png"
    if _sig_png.exists():
        from PIL import Image as _PILImage
        with _PILImage.open(str(_sig_png)) as _im:
            _iw, _ih = _im.size
        _sig_w = min(doc.width, _AVAIL_W)
        _sig_h = _ih * (_sig_w / _iw)
        story.append(KeepTogether([
            Spacer(1, _SPACER_H),
            RLImage(str(_sig_png), width=_sig_w, height=_sig_h),
        ]))

    doc.build(story, canvasmaker=_make_inv_canvas(merged_no, date_ru, date_en))

    await message.answer_document(
        FSInputFile(str(inv_pdf)),
        caption="✅ PDF инвойс",
    )


@router.callback_query(F.data == "merge:restart")
async def on_merge_restart(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    user_id = callback.from_user.id
    tmp_dir = TMP_DIR / str(user_id) / "specs"
    if tmp_dir.exists():
        shutil.rmtree(tmp_dir)
    tmp_dir.mkdir(parents=True, exist_ok=True)
    await state.set_state(MergeSpec.waiting_pdfs)
    await state.update_data(spec_files=[], spec_infos=[], pdf_files=[], status_msg_id=None)
    await callback.message.answer(
        "🔗 Новое объединение\n\n"
        "Загрузите <b>XLSX-файлы</b> спецификаций и <b>один PDF</b> "
        "(с блоком печатей и подписей, не разделённым между страницами).\n\n"
        "Дождитесь, пока система примет все файлы.",
        parse_mode="HTML",
        reply_markup=_merge_upload_kb(),
    )
